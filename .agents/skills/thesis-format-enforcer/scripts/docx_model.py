#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
ANCHOR_LABELS = {
    "abstract_cn",
    "abstract_en",
    "toc",
    "references",
    "chapter_heading",
    "section_heading",
}
NOTE_KEYWORDS = ("要求", "说明", "注", "格式", "字号", "字体", "行距", "页边距")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_docx_part(path: Path, part_name: str) -> etree._Element | None:
    with zipfile.ZipFile(path) as archive:
        try:
            data = archive.read(part_name)
        except KeyError:
            return None
    return etree.fromstring(data)


def safe_text(value: Any) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ").replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def twips_to_pt(value: str | None) -> float | None:
    if value is None:
        return None
    try:
        return round(int(value) / 20.0, 2)
    except (TypeError, ValueError):
        return None


def half_points_to_pt(value: str | None) -> float | None:
    if value is None:
        return None
    try:
        return round(int(value) / 2.0, 2)
    except (TypeError, ValueError):
        return None


def length_to_pt(value: Any) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        return round(float(value.pt), 2)
    if isinstance(value, (int, float)):
        return round(float(value), 2)
    return None


def normalize_bool(value: Any) -> bool | None:
    if value is None:
        return None
    return bool(value)


def normalize_alignment(value: Any) -> str | None:
    if value is None:
        return None
    if hasattr(value, "name") and value.name:
        return value.name.lower()
    if isinstance(value, str):
        return value.lower()
    return safe_text(value)


def xml_on_off(element: etree._Element | None) -> bool | None:
    if element is None:
        return None
    value = element.get(qn("w:val"))
    if value is None:
        return True
    return value not in {"0", "false", "off"}


def first_child(parent: etree._Element | None, tag: str) -> etree._Element | None:
    if parent is None:
        return None
    return parent.find(f"w:{tag}", NS)


def child_val(parent: etree._Element | None, tag: str) -> str | None:
    child = first_child(parent, tag)
    if child is None:
        return None
    return child.get(qn("w:val"))


def extract_run_props_xml(rpr: etree._Element | None) -> dict[str, Any]:
    if rpr is None:
        return {}
    fonts = first_child(rpr, "rFonts")
    return {
        "font_name": safe_text(fonts.get(qn("w:ascii")) if fonts is not None else None),
        "east_asia_font": safe_text(fonts.get(qn("w:eastAsia")) if fonts is not None else None),
        "size_pt": half_points_to_pt(child_val(rpr, "sz")),
        "bold": xml_on_off(first_child(rpr, "b")),
        "italic": xml_on_off(first_child(rpr, "i")),
        "color": safe_text(child_val(rpr, "color")),
    }


def extract_paragraph_props_xml(ppr: etree._Element | None) -> dict[str, Any]:
    if ppr is None:
        return {}
    spacing = first_child(ppr, "spacing")
    indent = first_child(ppr, "ind")
    jc = child_val(ppr, "jc")
    return {
        "alignment": safe_text(jc.lower() if jc else None),
        "space_before_pt": twips_to_pt(spacing.get(qn("w:before")) if spacing is not None else None),
        "space_after_pt": twips_to_pt(spacing.get(qn("w:after")) if spacing is not None else None),
        "line_spacing_pt": twips_to_pt(spacing.get(qn("w:line")) if spacing is not None else None),
        "line_spacing_rule": safe_text(spacing.get(qn("w:lineRule")) if spacing is not None else None),
        "left_indent_pt": twips_to_pt(indent.get(qn("w:left")) if indent is not None else None),
        "right_indent_pt": twips_to_pt(indent.get(qn("w:right")) if indent is not None else None),
        "first_line_indent_pt": twips_to_pt(
            indent.get(qn("w:firstLine")) if indent is not None else None
        ),
        "hanging_indent_pt": twips_to_pt(
            indent.get(qn("w:hanging")) if indent is not None else None
        ),
        "keep_next": xml_on_off(first_child(ppr, "keepNext")),
        "keep_lines": xml_on_off(first_child(ppr, "keepLines")),
        "page_break_before": xml_on_off(first_child(ppr, "pageBreakBefore")),
        "widow_control": xml_on_off(first_child(ppr, "widowControl")),
        "outline_level": safe_text(child_val(ppr, "outlineLvl")),
    }


def merge_specs(*specs: dict[str, Any] | None) -> dict[str, Any]:
    merged: dict[str, Any] = {}
    for spec in specs:
        if not spec:
            continue
        for key, value in spec.items():
            if value is not None:
                merged[key] = value
    return merged


def paragraph_direct_props(paragraph: Any) -> dict[str, Any]:
    ppr = paragraph._p.pPr
    if ppr is None:
        return {}
    return extract_paragraph_props_xml(ppr)


def run_direct_props(run: Any) -> dict[str, Any]:
    rpr = run._r.rPr
    direct = {
        "font_name": safe_text(run.font.name),
        "size_pt": length_to_pt(run.font.size),
        "bold": normalize_bool(run.font.bold),
        "italic": normalize_bool(run.font.italic),
    }
    return merge_specs(extract_run_props_xml(rpr), direct)


def style_type_name(style: Any) -> str | None:
    value = getattr(style, "type", None)
    if value is None:
        return None
    if hasattr(value, "name") and value.name:
        return value.name.lower()
    return safe_text(value)


def extract_style_specs(document: Document, styles_root: etree._Element | None) -> dict[str, Any]:
    if styles_root is None:
        return {}
    style_meta: dict[str, dict[str, Any]] = {}
    for style in styles_root.findall("w:style", NS):
        style_id = safe_text(style.get(qn("w:styleId")))
        if not style_id:
            continue
        style_meta[style_id] = {
            "style_id": style_id,
            "name": safe_text(child_val(style, "name")),
            "based_on": safe_text(child_val(style, "basedOn")),
            "next_style": safe_text(child_val(style, "next")),
            "style_type": safe_text(style.get(qn("w:type"))),
            "run": extract_run_props_xml(first_child(style, "rPr")),
            "paragraph": extract_paragraph_props_xml(first_child(style, "pPr")),
        }

    resolved: dict[str, Any] = {}
    for style in document.styles:
        style_id = safe_text(getattr(style, "style_id", None))
        if not style_id:
            continue
        meta = style_meta.get(style_id, {})
        resolved[style_id] = {
            "style_id": style_id,
            "name": safe_text(getattr(style, "name", None)) or meta.get("name"),
            "style_type": style_type_name(style) or meta.get("style_type"),
            "based_on": meta.get("based_on"),
            "next_style": meta.get("next_style"),
            "run": meta.get("run", {}),
            "paragraph": meta.get("paragraph", {}),
        }
    return resolved


def extract_doc_defaults(styles_root: etree._Element | None) -> dict[str, Any]:
    if styles_root is None:
        return {"run": {}, "paragraph": {}}
    doc_defaults = first_child(styles_root, "docDefaults")
    run_default = first_child(first_child(doc_defaults, "rPrDefault"), "rPr")
    para_default = first_child(first_child(doc_defaults, "pPrDefault"), "pPr")
    return {
        "run": extract_run_props_xml(run_default),
        "paragraph": extract_paragraph_props_xml(para_default),
    }


def section_snapshot(section: Any, index: int) -> dict[str, Any]:
    start_type = getattr(section.start_type, "name", None)
    orientation = getattr(section.orientation, "name", None)
    return {
        "index": index,
        "start_type": safe_text(start_type.lower() if start_type else None),
        "orientation": safe_text(orientation.lower() if orientation else None),
        "page_width_pt": length_to_pt(section.page_width),
        "page_height_pt": length_to_pt(section.page_height),
        "left_margin_pt": length_to_pt(section.left_margin),
        "right_margin_pt": length_to_pt(section.right_margin),
        "top_margin_pt": length_to_pt(section.top_margin),
        "bottom_margin_pt": length_to_pt(section.bottom_margin),
        "header_distance_pt": length_to_pt(section.header_distance),
        "footer_distance_pt": length_to_pt(section.footer_distance),
        "gutter_pt": length_to_pt(getattr(section, "gutter", None)),
    }


def extract_section_blueprint(document_root: etree._Element | None) -> dict[str, Any]:
    if document_root is None:
        return {"section_breaks": [], "final_section_xml": None}
    body = document_root.find("w:body", NS)
    if body is None:
        return {"section_breaks": [], "final_section_xml": None}

    section_breaks: list[dict[str, Any]] = []
    paragraph_index = 0
    for child in body:
        if child.tag != qn("w:p"):
            continue
        sect_pr = child.find("w:pPr/w:sectPr", NS)
        if sect_pr is not None:
            section_breaks.append(
                {
                    "after_paragraph_index": paragraph_index,
                    "sectPr_xml": etree.tostring(sect_pr, encoding="unicode"),
                }
            )
        paragraph_index += 1

    final_section = body.find("w:sectPr", NS)
    return {
        "section_breaks": section_breaks,
        "final_section_xml": etree.tostring(final_section, encoding="unicode") if final_section is not None else None,
    }


def representative_run(paragraph: Any) -> Any | None:
    chosen = None
    chosen_weight = -1
    for run in paragraph.runs:
        weight = len(normalize_text(run.text))
        if weight > chosen_weight:
            chosen = run
            chosen_weight = weight
    return chosen


def infer_semantic_label(text: str, style_name: str | None, outline_level: str | None) -> str | None:
    if not text:
        return None
    squashed = re.sub(r"\s+", "", text).lower()
    style_name = (style_name or "").lower()
    if "toc" in style_name or "目录" in style_name:
        if squashed not in {"目录", "contents", "tableofcontents"}:
            return None
    if squashed == "摘要":
        return "abstract_cn"
    if squashed == "abstract":
        return "abstract_en"
    if squashed in {"关键词", "关键字"}:
        return "keywords_cn"
    if squashed in {"keywords", "keywords:"}:
        return "keywords_en"
    if squashed in {"目录", "contents", "tableofcontents"}:
        return "toc"
    if squashed in {"参考文献", "references", "bibliography"}:
        return "references"
    if squashed in {"致谢", "acknowledgements", "acknowledgments"}:
        return "acknowledgements"
    if squashed.startswith("附录") or squashed.startswith("appendix"):
        return "appendix"
    if re.match(r"^第[一二三四五六七八九十百零0-9]+章", squashed):
        return "chapter_heading"
    if re.match(r"^chapter[0-9ivxlcdm]+$", squashed):
        return "chapter_heading"
    if re.match(r"^[0-9]+(\.[0-9]+)+", squashed):
        return "section_heading"
    if re.match(r"^第[一二三四五六七八九十百零0-9]+节", squashed):
        return "section_heading"
    if outline_level == "0" or "heading 1" in style_name or "标题 1" in style_name:
        return "chapter_heading"
    if outline_level == "1" or "heading 2" in style_name or "标题 2" in style_name:
        return "section_heading"
    return None


def paragraph_snapshot(paragraph: Any) -> dict[str, Any]:
    return {
        "index": paragraph["index"],
        "text": paragraph["text"],
        "style_name": paragraph["style_name"],
        "style_id": paragraph["style_id"],
        "semantic_label": paragraph["semantic_label"],
        "paragraph": paragraph["paragraph"],
        "run": paragraph["run"],
        "runs": paragraph.get("runs", []),
        "is_empty": paragraph.get("is_empty", False),
    }


def run_snapshot(
    run: Any,
    *,
    defaults: dict[str, Any],
    paragraph_style_spec: dict[str, Any] | None,
    style_specs: dict[str, Any],
    style_by_name: dict[str, Any],
) -> dict[str, Any]:
    style_name = safe_text(getattr(run.style, "name", None))
    style_id = safe_text(getattr(run.style, "style_id", None))
    style_spec = style_specs.get(style_id) or style_by_name.get(style_name)
    effective_run = merge_specs(
        defaults.get("run"),
        paragraph_style_spec.get("run") if paragraph_style_spec else None,
        style_spec.get("run") if style_spec else None,
        run_direct_props(run),
    )
    return {
        "text": run.text or "",
        "style_name": style_name,
        "style_id": style_id,
        "is_empty": not bool(run.text),
        "run": effective_run,
    }


def paragraph_record(
    paragraph: Any,
    *,
    index: int,
    style_specs: dict[str, Any],
    defaults: dict[str, Any],
) -> dict[str, Any]:
    style_by_name = {
        spec["name"]: spec for spec in style_specs.values() if spec.get("name")
    }
    style_name = safe_text(getattr(paragraph.style, "name", None))
    style_id = safe_text(getattr(paragraph.style, "style_id", None))
    style_spec = style_specs.get(style_id) or style_by_name.get(style_name)
    direct_para = paragraph_direct_props(paragraph)
    effective_para = merge_specs(
        defaults.get("paragraph"),
        style_spec.get("paragraph") if style_spec else None,
        direct_para,
    )
    run_snapshots = [
        run_snapshot(
            run,
            defaults=defaults,
            paragraph_style_spec=style_spec,
            style_specs=style_specs,
            style_by_name=style_by_name,
        )
        for run in paragraph.runs
    ]
    representative = max(
        run_snapshots,
        key=lambda item: len(normalize_text(item["text"])),
        default=None,
    )
    effective_run = representative["run"] if representative is not None else {}
    text = normalize_text(paragraph.text)
    semantic_label = infer_semantic_label(
        text=text,
        style_name=style_name,
        outline_level=effective_para.get("outline_level"),
    )
    return {
        "index": index,
        "text": text,
        "style_name": style_name,
        "style_id": style_id,
        "is_empty": not bool(text),
        "semantic_label": semantic_label,
        "paragraph": effective_para,
        "run": effective_run,
        "runs": run_snapshots,
    }


def extract_paragraphs(
    document: Document,
    style_specs: dict[str, Any],
    defaults: dict[str, Any],
) -> list[dict[str, Any]]:
    return [
        paragraph_record(
            paragraph,
            index=index,
            style_specs=style_specs,
            defaults=defaults,
        )
        for index, paragraph in enumerate(document.paragraphs)
    ]


def extract_header_paragraphs(
    document: Document,
    style_specs: dict[str, Any],
    defaults: dict[str, Any],
) -> list[dict[str, Any]]:
    header_paragraphs: list[dict[str, Any]] = []
    header_sources = [
        ("default", "header"),
        ("first_page", "first_page_header"),
        ("even_page", "even_page_header"),
    ]
    index = 0
    for section_index, section in enumerate(document.sections):
        for header_kind, attr_name in header_sources:
            header = getattr(section, attr_name, None)
            if header is None:
                continue
            for header_paragraph_index, paragraph in enumerate(header.paragraphs):
                item = paragraph_record(
                    paragraph,
                    index=index,
                    style_specs=style_specs,
                    defaults=defaults,
                )
                item["semantic_label"] = None
                item["section_index"] = section_index
                item["header_kind"] = header_kind
                item["header_paragraph_index"] = header_paragraph_index
                header_paragraphs.append(item)
                index += 1
    return header_paragraphs


def extract_table_paragraphs(
    document: Document,
    style_specs: dict[str, Any],
    defaults: dict[str, Any],
) -> list[dict[str, Any]]:
    table_paragraphs: list[dict[str, Any]] = []
    index = 0
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                for cell_paragraph_index, paragraph in enumerate(cell.paragraphs):
                    item = paragraph_record(
                        paragraph,
                        index=index,
                        style_specs=style_specs,
                        defaults=defaults,
                    )
                    item["semantic_label"] = None
                    item["table_index"] = table_index
                    item["row_index"] = row_index
                    item["column_index"] = column_index
                    item["cell_paragraph_index"] = cell_paragraph_index
                    table_paragraphs.append(item)
                    index += 1
    return table_paragraphs


def collect_required_sections(paragraphs: list[dict[str, Any]]) -> list[str]:
    ordered: list[str] = []
    for paragraph in paragraphs:
        label = paragraph.get("semantic_label")
        if label and label not in ordered and label not in {"keywords_cn", "keywords_en"}:
            ordered.append(label)
    return ordered


def collect_cover_blocks(paragraphs: list[dict[str, Any]], limit: int = 20) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    blank_count = 0
    for paragraph in paragraphs:
        if paragraph["semantic_label"] in ANCHOR_LABELS and blocks:
            break
        if paragraph["is_empty"]:
            blank_count += 1
            continue
        block = paragraph_snapshot(paragraph)
        block["leading_blank_paragraphs"] = blank_count
        blocks.append(block)
        blank_count = 0
        if len(blocks) >= limit:
            break
    return blocks


def collect_semantic_rules(paragraphs: list[dict[str, Any]]) -> dict[str, Any]:
    rules: dict[str, Any] = {}
    for paragraph in paragraphs:
        label = paragraph.get("semantic_label")
        if label and label not in rules:
            rules[label] = paragraph_snapshot(paragraph)
    return rules


def collect_heading_rules(paragraphs: list[dict[str, Any]]) -> dict[str, Any]:
    rules: dict[str, Any] = {}
    for paragraph in paragraphs:
        key = paragraph["paragraph"].get("outline_level")
        if key is None:
            label = paragraph.get("semantic_label")
            if label == "chapter_heading":
                key = "0"
            elif label == "section_heading":
                key = "1"
        if key is None:
            continue
        if key not in rules:
            rules[key] = paragraph_snapshot(paragraph)
    return rules


def collect_requirement_notes(paragraphs: list[dict[str, Any]], limit: int = 40) -> list[dict[str, Any]]:
    notes: list[dict[str, Any]] = []
    for paragraph in paragraphs:
        text = paragraph["text"]
        if not text or len(text) > 120:
            continue
        if any(keyword in text for keyword in NOTE_KEYWORDS):
            notes.append({"index": paragraph["index"], "text": text})
        if len(notes) >= limit:
            break
    return notes


def collect_tracked_styles(
    style_specs: dict[str, Any],
    paragraphs: list[dict[str, Any]],
    headers: list[dict[str, Any]],
    table_paragraphs: list[dict[str, Any]],
    semantic_rules: dict[str, Any],
    cover_blocks: list[dict[str, Any]],
    heading_rules: dict[str, Any],
) -> list[dict[str, Any]]:
    wanted_ids: set[str] = set()
    wanted_names: set[str] = {"Normal"}
    for paragraph in [*paragraphs, *headers, *table_paragraphs]:
        style_id = paragraph.get("style_id")
        style_name = paragraph.get("style_name")
        if style_id:
            wanted_ids.add(style_id)
        if style_name:
            wanted_names.add(style_name)
        for run in paragraph.get("runs", []):
            run_style_id = run.get("style_id")
            run_style_name = run.get("style_name")
            if run_style_id:
                wanted_ids.add(run_style_id)
            if run_style_name:
                wanted_names.add(run_style_name)
    for source in [*semantic_rules.values(), *cover_blocks, *heading_rules.values()]:
        style_id = source.get("style_id")
        style_name = source.get("style_name")
        if style_id:
            wanted_ids.add(style_id)
        if style_name:
            wanted_names.add(style_name)
    tracked = []
    for style in style_specs.values():
        if style["style_id"] in wanted_ids or style.get("name") in wanted_names:
            tracked.append(style)
            continue
        if style.get("paragraph", {}).get("outline_level") is not None:
            tracked.append(style)
    tracked.sort(key=lambda item: (item.get("style_type") or "", item.get("name") or ""))
    return tracked


def render_paragraph_index(paragraphs: list[dict[str, Any]], limit: int = 120) -> str:
    lines = [
        "# Template Paragraph Index",
        "",
        "| # | Style | Label | Text |",
        "| --- | --- | --- | --- |",
    ]
    for paragraph in paragraphs[:limit]:
        text = paragraph["text"].replace("|", "\\|")
        lines.append(
            "| {index} | {style} | {label} | {text} |".format(
                index=paragraph["index"],
                style=paragraph.get("style_name") or "",
                label=paragraph.get("semantic_label") or "",
                text=text[:120],
            )
        )
    if len(paragraphs) > limit:
        lines.extend(["", f"Showing first {limit} paragraphs out of {len(paragraphs)}."])
    lines.append("")
    return "\n".join(lines)


def build_document_model(document_path: Path) -> dict[str, Any]:
    document = Document(str(document_path))
    styles_root = read_docx_part(document_path, "word/styles.xml")
    document_root = read_docx_part(document_path, "word/document.xml")
    defaults = extract_doc_defaults(styles_root)
    style_specs = extract_style_specs(document, styles_root)
    paragraphs = extract_paragraphs(document, style_specs, defaults)
    headers = extract_header_paragraphs(document, style_specs, defaults)
    table_paragraphs = extract_table_paragraphs(document, style_specs, defaults)
    semantic_rules = collect_semantic_rules(paragraphs)
    cover_blocks = collect_cover_blocks(paragraphs)
    heading_rules = collect_heading_rules(paragraphs)
    return {
        "source": {
            "path": str(document_path.resolve()),
            "file_name": document_path.name,
            "sha256": sha256_file(document_path),
        },
        "document_defaults": defaults,
        "styles": style_specs,
        "sections": [section_snapshot(section, index) for index, section in enumerate(document.sections)],
        "section_blueprint": extract_section_blueprint(document_root),
        "paragraphs": paragraphs,
        "headers": headers,
        "table_paragraphs": table_paragraphs,
        "required_sections": collect_required_sections(paragraphs),
        "semantic_rules": semantic_rules,
        "cover_blocks": cover_blocks,
        "heading_rules": heading_rules,
        "requirement_notes": collect_requirement_notes(paragraphs),
        "paragraph_index_markdown": render_paragraph_index(paragraphs),
        "tracked_styles": collect_tracked_styles(
            style_specs=style_specs,
            paragraphs=paragraphs,
            headers=headers,
            table_paragraphs=table_paragraphs,
            semantic_rules=semantic_rules,
            cover_blocks=cover_blocks,
            heading_rules=heading_rules,
        ),
    }
