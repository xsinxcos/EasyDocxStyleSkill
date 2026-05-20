#!/usr/bin/env python3
from __future__ import annotations

import argparse
import copy
import hashlib
import json
import os
import re
import shutil
import sys
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from docx_model import build_document_model

PIPELINE_VERSION = "0.1.0"
RULE_PACK_VERSION = 1
REGISTRY_VERSION = 1

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}
ORIENTATION_MAP = {
    "portrait": WD_ORIENT.PORTRAIT,
    "landscape": WD_ORIENT.LANDSCAPE,
}
STYLE_TYPE_MAP = {
    "paragraph": WD_STYLE_TYPE.PARAGRAPH,
    "character": WD_STYLE_TYPE.CHARACTER,
    "table": WD_STYLE_TYPE.TABLE,
}
SIZE_LABEL_MAP = {
    "初号": 42.0,
    "小初": 36.0,
    "一号": 26.0,
    "小一": 24.0,
    "二号": 22.0,
    "小二": 18.0,
    "三号": 16.0,
    "小三": 15.0,
    "四号": 14.0,
    "小四": 12.0,
    "五号": 10.5,
    "小五": 9.0,
}
TEXT_REQUIREMENT_LABELS = {
    "abstract_cn_body": "中文摘要正文",
    "abstract_en_body": "英文摘要正文",
    "body_paragraph": "正文段落",
    "reference_entry": "参考文献条目",
    "heading_level_1": "一级标题要求",
    "heading_level_2": "二级标题要求",
    "heading_level_3": "三级标题要求",
    "table_caption": "表题要求",
    "table_text": "表格正文要求",
    "figure_caption": "图题要求",
    "code_block": "代码块要求",
    "header_text": "页眉要求",
    "global_consistency": "全文一致性要求",
    "caption_consistency": "图表公式标注一致性要求",
}


def utc_now() -> str:
    return datetime.now(UTC).isoformat()


def workspace_root(start: Path) -> Path:
    for candidate in [start, *start.parents]:
        if (candidate / ".git").exists() or (candidate / ".codex").exists():
            return candidate
    return start


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def relative_path(path: Path, root: Path) -> str:
    try:
        return str(path.resolve().relative_to(root.resolve()))
    except ValueError:
        return str(path.resolve())


def slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9._-]+", "-", value.strip().lower())
    slug = re.sub(r"-{2,}", "-", slug).strip("-")
    return slug or "rule-pack"


def path_slug(path: Path, root: Path) -> str:
    try:
        relative = path.resolve().relative_to(root.resolve())
        base = "-".join(relative.with_suffix("").parts)
    except ValueError:
        base = path.stem
    return slugify(base)


def read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def write_json(path: Path, payload: Any) -> None:
    ensure_dir(path.parent)
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2)
        handle.write("\n")


def write_text(path: Path, payload: str) -> None:
    ensure_dir(path.parent)
    path.write_text(payload, encoding="utf-8")


def registry_path(root: Path) -> Path:
    return root / "artifacts" / ".workflow" / "thesis-format-registry.json"


def load_registry(root: Path) -> dict[str, Any]:
    data = read_json(registry_path(root), default={})
    if not data:
        return {
            "version": REGISTRY_VERSION,
            "rule_packs": [],
            "audits": [],
            "repairs": [],
            "latest_rule_pack": None,
            "latest_audit": None,
            "latest_repair": None,
        }
    return data


def save_registry(root: Path, registry: dict[str, Any]) -> None:
    write_json(registry_path(root), registry)


def find_rule_pack_path(root: Path, value: str | None) -> Path | None:
    if value:
        candidate = Path(value)
        if not candidate.is_absolute():
            candidate = (root / candidate).resolve()
        if (candidate / "manifest.json").exists() and (candidate / "rules.json").exists():
            return candidate
        return candidate if candidate.exists() else None
    registry = load_registry(root)
    latest = registry.get("latest_rule_pack")
    if not latest:
        return None
    candidate = Path(latest["path"])
    return candidate if candidate.exists() else None


def find_matching_audit(
    registry: dict[str, Any],
    document_path: Path,
    rule_pack_path: Path,
) -> dict[str, Any] | None:
    document_path = document_path.resolve()
    rule_pack_path = rule_pack_path.resolve()
    for audit in reversed(registry.get("audits", [])):
        if (
            Path(audit["document_path"]).resolve() == document_path
            and Path(audit["rule_pack_path"]).resolve() == rule_pack_path
        ):
            return audit
    return None


def load_rule_pack(rule_pack_path: Path) -> dict[str, Any]:
    manifest = read_json(rule_pack_path / "manifest.json", default=None)
    rules = read_json(rule_pack_path / "rules.json", default=None)
    if manifest is None or rules is None:
        raise FileNotFoundError(f"Invalid rule pack: {rule_pack_path}")
    return {"manifest": manifest, "rules": rules}


def validation_artifact_dir(rule_pack_path: Path) -> Path:
    return ensure_dir(rule_pack_path / "validation")


def package_snapshot_dir(rule_pack_path: Path) -> Path:
    return rule_pack_path / "package-snapshot"


def package_manifest_path(rule_pack_path: Path) -> Path:
    return rule_pack_path / "package-manifest.json"


def sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def reset_dir(path: Path) -> Path:
    if path.exists():
        shutil.rmtree(path, onerror=handle_remove_readonly)
    path.mkdir(parents=True, exist_ok=True)
    return path


def handle_remove_readonly(func: Any, path: str, exc_info: Any) -> None:
    os.chmod(path, 0o666)
    func(path)


def extract_docx_package_snapshot(docx_path: Path, output_dir: Path) -> dict[str, Any]:
    ensure_dir(output_dir)
    entries: list[dict[str, Any]] = []
    with zipfile.ZipFile(docx_path) as archive:
        for name in sorted(info.filename for info in archive.infolist() if not info.is_dir()):
            payload = archive.read(name)
            target_path = output_dir / Path(*name.split("/"))
            ensure_dir(target_path.parent)
            target_path.write_bytes(payload)
            entries.append(
                {
                    "name": name,
                    "sha256": sha256_bytes(payload),
                    "size": len(payload),
                }
            )
    return {
        "entry_count": len(entries),
        "entries": entries,
    }


def manifest_from_docx_package(docx_path: Path) -> dict[str, Any]:
    entries: list[dict[str, Any]] = []
    with zipfile.ZipFile(docx_path) as archive:
        for name in sorted(info.filename for info in archive.infolist() if not info.is_dir()):
            payload = archive.read(name)
            entries.append(
                {
                    "name": name,
                    "sha256": sha256_bytes(payload),
                    "size": len(payload),
                }
            )
    return {
        "entry_count": len(entries),
        "entries": entries,
    }


def rebuild_docx_from_package_snapshot(snapshot_dir: Path, output_path: Path) -> None:
    if not snapshot_dir.exists():
        raise FileNotFoundError(f"Package snapshot not found: {snapshot_dir}")
    ensure_dir(output_path.parent)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(snapshot_dir.rglob("*")):
            if path.is_dir():
                continue
            arcname = path.relative_to(snapshot_dir).as_posix()
            archive.write(path, arcname)


def compare_package_manifests(expected: dict[str, Any], actual: dict[str, Any]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    expected_entries = {entry["name"]: entry for entry in expected.get("entries", [])}
    actual_entries = {entry["name"]: entry for entry in actual.get("entries", [])}

    for name in sorted(expected_entries):
        if name not in actual_entries:
            add_finding(
                findings,
                severity="error",
                category="docx-package",
                location=name,
                field="exists",
                expected="present",
                actual="missing",
                message="DOCX package part is missing from the rebuilt template.",
            )
            continue
        compare_mapping(
            findings,
            expected=expected_entries[name],
            actual=actual_entries[name],
            category="docx-package",
            location=name,
            severity="error",
            fields=["sha256", "size"],
        )

    for name in sorted(actual_entries):
        if name in expected_entries:
            continue
        add_finding(
            findings,
            severity="error",
            category="docx-package",
            location=name,
            field="exists",
            expected="missing",
            actual="present",
            message="DOCX package part is extra in the rebuilt template.",
        )
    return findings


def template_blueprint_from_model(model: dict[str, Any]) -> dict[str, Any]:
    return {
        "paragraphs": [
            {
                "index": paragraph["index"],
                "text": paragraph["text"],
                "style_name": paragraph.get("style_name"),
                "style_id": paragraph.get("style_id"),
                "semantic_label": paragraph.get("semantic_label"),
                "paragraph": paragraph.get("paragraph", {}),
                "run": paragraph.get("run", {}),
                "runs": paragraph.get("runs", []),
                "is_empty": paragraph.get("is_empty", False),
            }
            for paragraph in model.get("paragraphs", [])
        ],
        "paragraph_count": len(model.get("paragraphs", [])),
        "section_count": len(model.get("sections", [])),
        "section_breaks": model.get("section_blueprint", {}).get("section_breaks", []),
        "final_section_xml": model.get("section_blueprint", {}).get("final_section_xml"),
    }


def header_rules_from_model(model: dict[str, Any]) -> dict[str, Any]:
    rules: dict[str, Any] = {}
    for header in model.get("headers", []):
        header_kind = header.get("header_kind")
        if not header_kind or header.get("is_empty"):
            continue
        if header_kind in rules:
            continue
        rules[header_kind] = {
            "header_kind": header_kind,
            "style_name": header.get("style_name"),
            "paragraph": copy.deepcopy(header.get("paragraph", {})),
            "run": copy.deepcopy(header.get("run", {})),
            "text": header.get("text"),
            "source_section_index": header.get("section_index"),
            "source_header_paragraph_index": header.get("header_paragraph_index"),
            "constraints": {},
        }
    return rules


def find_style_rule(rules: dict[str, Any], style_name: str) -> dict[str, Any] | None:
    for style in rules.get("tracked_styles", []):
        if style.get("name") == style_name:
            return style
    return None


def merge_non_none(target: dict[str, Any], override: dict[str, Any]) -> None:
    for key, value in override.items():
        if value is not None:
            target[key] = value


def line_count_to_points(value: float | None, line_spacing_pt: float | None) -> float | None:
    if value is None or line_spacing_pt is None:
        return None
    return round(value * line_spacing_pt, 2)


def extract_size_from_text(text: str) -> float | None:
    for label, value in sorted(SIZE_LABEL_MAP.items(), key=lambda item: len(item[0]), reverse=True):
        if label in text:
            return value
    return None


def extract_requirement_spec(note_text: str, target: str) -> dict[str, Any]:
    working_text = note_text
    if target == "abstract_cn_body" and "“摘要”" in working_text:
        working_text = working_text.split("“摘要”", 1)[0]

    run_spec: dict[str, Any] = {}
    paragraph_spec: dict[str, Any] = {}
    constraints: dict[str, Any] = {}

    chinese_fonts = [font for font in ["宋体", "黑体", "仿宋", "楷体"] if font in working_text]
    english_fonts = [font for font in ["Times New Roman", "Arial", "Consolas"] if font in working_text]
    size_pt = extract_size_from_text(working_text)
    line_spacing_match = re.search(r"(?:行距(?:固定值)?|固定值)\s*([0-9]+(?:\.[0-9]+)?)\s*磅", working_text)
    line_spacing_pt = round(float(line_spacing_match.group(1)), 2) if line_spacing_match else None
    space_before_lines_match = re.search(r"段前\s*([0-9]+(?:\.[0-9]+)?)\s*行", working_text)
    space_after_lines_match = re.search(r"段后\s*([0-9]+(?:\.[0-9]+)?)\s*行", working_text)
    first_indent_match = re.search(r"首行缩进\s*([0-9]+(?:\.[0-9]+)?)\s*(?:个)?字符", working_text)
    recommended_count_match = re.search(r"建议\s*([0-9]+)\s*-\s*([0-9]+)\s*条", working_text)
    word_limit_match = re.search(r"([0-9]+)\s*字以内", working_text)

    if chinese_fonts:
        run_spec["east_asia_font"] = chinese_fonts[0]
        if not english_fonts:
            run_spec["font_name"] = chinese_fonts[0]
    if english_fonts:
        run_spec["font_name"] = english_fonts[0]
    if size_pt is not None:
        run_spec["size_pt"] = size_pt
    if "加粗" in working_text:
        run_spec["bold"] = True
    if "斜体" in working_text:
        run_spec["italic"] = True

    if line_spacing_pt is not None:
        paragraph_spec["line_spacing_pt"] = line_spacing_pt
        paragraph_spec["line_spacing_rule"] = "exact"
    elif "单倍行距" in working_text:
        constraints["line_spacing_mode"] = "single"
    if "居中" in working_text:
        paragraph_spec["alignment"] = "center"
    if "两端对齐" in working_text:
        paragraph_spec["alignment"] = "both"
    if "左对齐" in working_text or "左顶格" in working_text:
        paragraph_spec["alignment"] = "left"
    if space_before_lines_match and line_spacing_pt is not None:
        paragraph_spec["space_before_pt"] = line_count_to_points(float(space_before_lines_match.group(1)), line_spacing_pt)
    if space_after_lines_match and line_spacing_pt is not None:
        paragraph_spec["space_after_pt"] = line_count_to_points(float(space_after_lines_match.group(1)), line_spacing_pt)
    if first_indent_match and size_pt is not None:
        paragraph_spec["first_line_indent_pt"] = round(float(first_indent_match.group(1)) * size_pt, 2)

    if "正文中一定要有引用" in working_text:
        constraints["citation_required_in_body"] = True
    if recommended_count_match:
        constraints["recommended_entry_count"] = [
            int(recommended_count_match.group(1)),
            int(recommended_count_match.group(2)),
        ]
    if word_limit_match:
        constraints["max_words"] = int(word_limit_match.group(1))
    numbering_match = re.search(r"编号为(.+)$", working_text)
    if numbering_match:
        constraints["numbering_pattern"] = numbering_match.group(1).strip("。； ")

    return {
        "run": run_spec,
        "paragraph": paragraph_spec,
        "constraints": constraints,
    }


def build_text_requirement_rules(model: dict[str, Any]) -> list[dict[str, Any]]:
    text_rules: list[dict[str, Any]] = []
    previous_by_target: dict[str, dict[str, Any]] = {}

    for note in model.get("requirement_notes", []):
        text = note["text"]
        target = None
        supported = False

        if text.startswith("中文摘要"):
            target = "abstract_cn_body"
            supported = True
        elif text.startswith("英文摘要"):
            target = "abstract_en_body"
            supported = True
        elif "一级标题格式" in text:
            target = "heading_level_1"
            supported = True
        elif "二级标题" in text and "三级标题" not in text:
            target = "heading_level_2"
            supported = True
        elif "三级标题" in text:
            target = "heading_level_3"
            supported = True
        elif text.startswith("正文格式"):
            target = "body_paragraph"
            supported = True
        elif "参考文献条目格式" in text:
            target = "reference_entry"
            supported = True
        elif text.startswith("所有的表都要有表号") or text.startswith("表题"):
            target = "table_caption"
        elif text.startswith("表中文本"):
            target = "table_text"
        elif text.startswith("所有的图都要有图号") or text.startswith("图题"):
            target = "figure_caption"
        elif "代码应该使用较小字体" in text or "Consolas字体" in text:
            target = "code_block"
        elif text.startswith("注意页眉"):
            target = "header_text"
        elif "图的标注、表的标注、公式标注" in text:
            target = "caption_consistency"
        elif "全文字体" in text and "统一风格" in text:
            target = "global_consistency"

        if target is None:
            continue

        spec = extract_requirement_spec(text, target)
        if target == "heading_level_3" and "同二级标题" in text:
            spec = copy.deepcopy(previous_by_target.get("heading_level_2", {}).get("spec", spec))
            numbering_match = re.search(r"编号为(.+)$", text)
            if numbering_match:
                spec.setdefault("constraints", {})["numbering_pattern"] = numbering_match.group(1).strip("。； ")

        applies_to: list[dict[str, Any]] = []
        if target == "heading_level_1":
            applies_to = [
                {"kind": "semantic_rule", "key": "chapter_heading"},
                {"kind": "tracked_style", "name": "Heading 1"},
            ]
        elif target == "heading_level_2":
            applies_to = [
                {"kind": "semantic_rule", "key": "section_heading"},
                {"kind": "tracked_style", "name": "Heading 2"},
            ]
        elif target == "heading_level_3":
            applies_to = [{"kind": "tracked_style", "name": "Heading 3"}]
        elif target == "body_paragraph":
            applies_to = [
                {"kind": "tracked_style", "name": "Normal"},
                {"kind": "custom_paragraph_rule", "name": "body_paragraph"},
            ]
        elif target == "abstract_cn_body":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "abstract_cn_body"}]
        elif target == "abstract_en_body":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "abstract_en_body"}]
        elif target == "reference_entry":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "reference_entry"}]

        requirement_rule = {
            "target": target,
            "label": TEXT_REQUIREMENT_LABELS.get(target, target),
            "source_paragraph_index": note["index"],
            "text": text,
            "supported": supported,
            "spec": spec,
            "applies_to": applies_to,
        }
        text_rules.append(requirement_rule)
        previous_by_target[target] = requirement_rule

    return text_rules


def build_custom_paragraph_rule(name: str, requirement_rule: dict[str, Any]) -> dict[str, Any]:
    scope: dict[str, Any]
    if name == "abstract_cn_body":
        scope = {
            "type": "between_labels",
            "start_label": "abstract_cn",
            "end_labels": ["keywords_cn", "abstract_en", "keywords_en", "toc"],
            "include_empty": False,
            "exclude_any_semantic_label": True,
            "allowed_style_names": ["Normal"],
        }
    elif name == "abstract_en_body":
        scope = {
            "type": "between_labels",
            "start_label": "abstract_en",
            "end_labels": ["keywords_en", "toc", "chapter_heading"],
            "include_empty": False,
            "exclude_any_semantic_label": True,
            "allowed_style_names": ["Normal"],
        }
    elif name == "body_paragraph":
        scope = {
            "type": "between_labels",
            "start_label": "chapter_heading",
            "end_labels": ["references", "acknowledgements", "appendix"],
            "include_empty": False,
            "exclude_any_semantic_label": True,
            "allowed_style_names": ["Normal"],
        }
    elif name == "reference_entry":
        scope = {
            "type": "between_labels",
            "start_label": "references",
            "end_labels": ["acknowledgements", "appendix"],
            "include_empty": False,
            "exclude_any_semantic_label": True,
            "text_regex": r"^\[\d+\]",
        }
    else:
        raise ValueError(f"Unsupported custom paragraph rule target: {name}")

    return {
        "name": name,
        "label": TEXT_REQUIREMENT_LABELS.get(name, name),
        "source_paragraph_index": requirement_rule["source_paragraph_index"],
        "text": requirement_rule["text"],
        "scope": scope,
        "paragraph": copy.deepcopy(requirement_rule["spec"].get("paragraph", {})),
        "run": copy.deepcopy(requirement_rule["spec"].get("run", {})),
        "constraints": copy.deepcopy(requirement_rule["spec"].get("constraints", {})),
    }


def apply_text_requirement_rules(rules: dict[str, Any], text_requirement_rules: list[dict[str, Any]]) -> dict[str, Any]:
    rules["text_requirement_rules"] = copy.deepcopy(text_requirement_rules)
    custom_paragraph_rules: list[dict[str, Any]] = []

    for requirement_rule in text_requirement_rules:
        if not requirement_rule.get("supported"):
            continue
        spec = requirement_rule.get("spec", {})
        paragraph_spec = spec.get("paragraph", {})
        run_spec = spec.get("run", {})

        for apply_target in requirement_rule.get("applies_to", []):
            kind = apply_target["kind"]
            if kind == "semantic_rule":
                target_rule = rules.get("semantic_rules", {}).get(apply_target["key"])
                if target_rule is None:
                    continue
                merge_non_none(target_rule.setdefault("paragraph", {}), paragraph_spec)
                merge_non_none(target_rule.setdefault("run", {}), run_spec)
            elif kind == "tracked_style":
                style_rule = find_style_rule(rules, apply_target["name"])
                if style_rule is None:
                    continue
                merge_non_none(style_rule.setdefault("paragraph", {}), paragraph_spec)
                merge_non_none(style_rule.setdefault("run", {}), run_spec)
            elif kind == "custom_paragraph_rule":
                custom_paragraph_rules.append(build_custom_paragraph_rule(apply_target["name"], requirement_rule))

    rules["custom_paragraph_rules"] = custom_paragraph_rules
    return rules


def build_text_requirement_rules_v2(model: dict[str, Any]) -> list[dict[str, Any]]:
    base_rules = build_text_requirement_rules(model)
    merged_rules: list[dict[str, Any]] = []
    by_target: dict[str, dict[str, Any]] = {}

    for base_rule in base_rules:
        target = base_rule.get("target")
        if not target:
            continue
        rule = copy.deepcopy(base_rule)
        rule["source_paragraph_indices"] = [base_rule.get("source_paragraph_index")]
        rule["texts"] = [base_rule.get("text")]

        if target == "table_caption":
            rule["supported"] = True
            rule.setdefault("spec", {}).setdefault("constraints", {})["text_regex"] = r"^(?:表|Table)\s*\d+[-－—–.．]\d+"
            rule["applies_to"] = [{"kind": "custom_paragraph_rule", "name": "table_caption"}]
        elif target == "figure_caption":
            rule["supported"] = True
            rule.setdefault("spec", {}).setdefault("constraints", {})["text_regex"] = r"^(?:图|Figure)\s*\d+[-－—–.．]\d+"
            rule["applies_to"] = [{"kind": "custom_paragraph_rule", "name": "figure_caption"}]
        elif target == "code_block":
            rule["supported"] = True
            rule["applies_to"] = [{"kind": "custom_paragraph_rule", "name": "code_block"}]
        elif target == "table_text":
            rule["supported"] = True
            rule["applies_to"] = [{"kind": "custom_table_rule", "name": "table_text"}]
        elif target == "header_text":
            rule["supported"] = True
            rule.setdefault("spec", {}).setdefault("constraints", {})["must_not_equal_expected_text"] = True
            rule["applies_to"] = [{"kind": "header_rule", "name": "default"}]

        existing = by_target.get(target)
        if existing is None:
            merged_rules.append(rule)
            by_target[target] = rule
            continue

        existing["supported"] = existing.get("supported") or rule.get("supported")
        existing.setdefault("source_paragraph_indices", []).extend(rule.get("source_paragraph_indices", []))
        existing.setdefault("texts", []).extend(rule.get("texts", []))
        existing["text"] = " / ".join(text for text in existing["texts"] if text)
        merge_non_none(existing.setdefault("spec", {}).setdefault("run", {}), rule.get("spec", {}).get("run", {}))
        merge_non_none(
            existing.setdefault("spec", {}).setdefault("paragraph", {}),
            rule.get("spec", {}).get("paragraph", {}),
        )
        merge_non_none(
            existing.setdefault("spec", {}).setdefault("constraints", {}),
            rule.get("spec", {}).get("constraints", {}),
        )
        for apply_target in rule.get("applies_to", []):
            if apply_target not in existing.setdefault("applies_to", []):
                existing["applies_to"].append(apply_target)

    return merged_rules

    text_rules: list[dict[str, Any]] = []
    previous_by_target: dict[str, dict[str, Any]] = {}
    rule_by_target: dict[str, dict[str, Any]] = {}

    for note in model.get("requirement_notes", []):
        text = note["text"]
        target = None
        supported = False

        if text.startswith("涓枃鎽樿"):
            target = "abstract_cn_body"
            supported = True
        elif text.startswith("鑻辨枃鎽樿"):
            target = "abstract_en_body"
            supported = True
        elif "涓€绾ф爣棰樻牸寮" in text:
            target = "heading_level_1"
            supported = True
        elif "浜岀骇鏍囬" in text and "涓夌骇鏍囬" not in text:
            target = "heading_level_2"
            supported = True
        elif "涓夌骇鏍囬" in text:
            target = "heading_level_3"
            supported = True
        elif text.startswith("姝ｆ枃鏍煎紡"):
            target = "body_paragraph"
            supported = True
        elif "鍙傝€冩枃鐚潯鐩牸寮" in text:
            target = "reference_entry"
            supported = True
        elif text.startswith("鎵€鏈夌殑琛ㄩ兘瑕佹湁琛ㄥ彿") or text.startswith("琛ㄩ"):
            target = "table_caption"
            supported = True
        elif text.startswith("琛ㄤ腑鏂囨湰"):
            target = "table_text"
            supported = True
        elif text.startswith("鎵€鏈夌殑鍥鹃兘瑕佹湁鍥惧彿") or text.startswith("鍥鹃"):
            target = "figure_caption"
            supported = True
        elif "浠ｇ爜搴旇浣跨敤杈冨皬瀛椾綋" in text or "Consolas瀛椾綋" in text:
            target = "code_block"
            supported = True
        elif text.startswith("娉ㄦ剰椤电湁"):
            target = "header_text"
            supported = True
        elif "鍥剧殑鏍囨敞銆佽〃鐨勬爣娉ㄣ€佸叕寮忔爣娉" in text:
            target = "caption_consistency"
        elif "鍏ㄦ枃瀛椾綋" in text and "缁熶竴椋庢牸" in text:
            target = "global_consistency"

        if target is None:
            continue

        spec = extract_requirement_spec(text, target)
        if target == "heading_level_3" and "鍚屼簩绾ф爣棰" in text:
            spec = copy.deepcopy(previous_by_target.get("heading_level_2", {}).get("spec", spec))
            numbering_match = re.search(r"缂栧彿涓?(.+)$", text)
            if numbering_match:
                spec.setdefault("constraints", {})["numbering_pattern"] = numbering_match.group(1).strip("銆傦紱 ")
        if target == "table_caption":
            spec.setdefault("constraints", {})["text_regex"] = r"^(?:表|Table)\s*\d+[-－—–.．]\d+"
        if target == "figure_caption":
            spec.setdefault("constraints", {})["text_regex"] = r"^(?:图|Figure)\s*\d+[-－—–.．]\d+"
        if target == "header_text":
            spec.setdefault("constraints", {})["must_not_equal_expected_text"] = True

        applies_to: list[dict[str, Any]] = []
        if target == "heading_level_1":
            applies_to = [
                {"kind": "semantic_rule", "key": "chapter_heading"},
                {"kind": "tracked_style", "name": "Heading 1"},
            ]
        elif target == "heading_level_2":
            applies_to = [
                {"kind": "semantic_rule", "key": "section_heading"},
                {"kind": "tracked_style", "name": "Heading 2"},
            ]
        elif target == "heading_level_3":
            applies_to = [{"kind": "tracked_style", "name": "Heading 3"}]
        elif target == "body_paragraph":
            applies_to = [
                {"kind": "tracked_style", "name": "Normal"},
                {"kind": "custom_paragraph_rule", "name": "body_paragraph"},
            ]
        elif target == "abstract_cn_body":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "abstract_cn_body"}]
        elif target == "abstract_en_body":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "abstract_en_body"}]
        elif target == "reference_entry":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "reference_entry"}]
        elif target == "table_caption":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "table_caption"}]
        elif target == "figure_caption":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "figure_caption"}]
        elif target == "code_block":
            applies_to = [{"kind": "custom_paragraph_rule", "name": "code_block"}]
        elif target == "table_text":
            applies_to = [{"kind": "custom_table_rule", "name": "table_text"}]
        elif target == "header_text":
            applies_to = [{"kind": "header_rule", "name": "default"}]

        existing = rule_by_target.get(target)
        if existing is None:
            requirement_rule = {
                "target": target,
                "label": TEXT_REQUIREMENT_LABELS.get(target, target),
                "source_paragraph_index": note["index"],
                "source_paragraph_indices": [note["index"]],
                "text": text,
                "texts": [text],
                "supported": supported,
                "spec": spec,
                "applies_to": applies_to,
            }
            text_rules.append(requirement_rule)
            rule_by_target[target] = requirement_rule
            existing = requirement_rule
        else:
            existing["supported"] = existing.get("supported") or supported
            existing.setdefault("source_paragraph_indices", []).append(note["index"])
            existing.setdefault("texts", []).append(text)
            existing["text"] = " / ".join(existing["texts"])
            merge_non_none(existing.setdefault("spec", {}).setdefault("run", {}), spec.get("run", {}))
            merge_non_none(existing.setdefault("spec", {}).setdefault("paragraph", {}), spec.get("paragraph", {}))
            merge_non_none(existing.setdefault("spec", {}).setdefault("constraints", {}), spec.get("constraints", {}))
            for apply_target in applies_to:
                if apply_target not in existing.setdefault("applies_to", []):
                    existing["applies_to"].append(apply_target)

        previous_by_target[target] = existing

    return text_rules


def build_custom_paragraph_rule_v2(name: str, requirement_rule: dict[str, Any]) -> dict[str, Any]:
    if name in {"abstract_cn_body", "abstract_en_body", "body_paragraph", "reference_entry"}:
        return build_custom_paragraph_rule(name, requirement_rule)

    if name == "table_caption":
        scope = {
            "type": "body_scope",
            "start_label": "chapter_heading",
            "end_labels": ["references", "acknowledgements", "appendix"],
            "include_empty": False,
            "exclude_any_semantic_label": True,
            "text_regex": r"^(?:表|Table)\s*\d+[-－—–.．]\d+",
        }
    elif name == "figure_caption":
        scope = {
            "type": "body_scope",
            "start_label": "chapter_heading",
            "end_labels": ["references", "acknowledgements", "appendix"],
            "include_empty": False,
            "exclude_any_semantic_label": True,
            "text_regex": r"^(?:图|Figure)\s*\d+[-－—–.．]\d+",
        }
    elif name == "code_block":
        scope = {
            "type": "body_scope",
            "start_label": "chapter_heading",
            "end_labels": ["references", "acknowledgements", "appendix"],
            "include_empty": False,
            "exclude_any_semantic_label": True,
            "heuristic": "code_like",
            "run_font_names": ["Consolas", "Courier New", "Lucida Console", "Monaco"],
            "max_text_length": 240,
        }
    else:
        raise ValueError(f"Unsupported custom paragraph rule target: {name}")

    return {
        "name": name,
        "label": TEXT_REQUIREMENT_LABELS.get(name, name),
        "source_paragraph_index": requirement_rule["source_paragraph_index"],
        "text": requirement_rule["text"],
        "scope": scope,
        "paragraph": copy.deepcopy(requirement_rule["spec"].get("paragraph", {})),
        "run": copy.deepcopy(requirement_rule["spec"].get("run", {})),
        "constraints": copy.deepcopy(requirement_rule["spec"].get("constraints", {})),
    }


def build_custom_table_rule_v2(name: str, requirement_rule: dict[str, Any]) -> dict[str, Any]:
    if name != "table_text":
        raise ValueError(f"Unsupported custom table rule target: {name}")
    return {
        "name": name,
        "label": TEXT_REQUIREMENT_LABELS.get(name, name),
        "source_paragraph_index": requirement_rule["source_paragraph_index"],
        "text": requirement_rule["text"],
        "scope": {
            "type": "all_table_paragraphs",
            "include_empty": False,
        },
        "paragraph": copy.deepcopy(requirement_rule["spec"].get("paragraph", {})),
        "run": copy.deepcopy(requirement_rule["spec"].get("run", {})),
        "constraints": copy.deepcopy(requirement_rule["spec"].get("constraints", {})),
    }


def apply_text_requirement_rules_v2(rules: dict[str, Any], text_requirement_rules: list[dict[str, Any]]) -> dict[str, Any]:
    rules["text_requirement_rules"] = copy.deepcopy(text_requirement_rules)
    custom_paragraph_rules: list[dict[str, Any]] = []
    custom_table_rules: list[dict[str, Any]] = []

    for requirement_rule in text_requirement_rules:
        if not requirement_rule.get("supported"):
            continue
        spec = requirement_rule.get("spec", {})
        paragraph_spec = spec.get("paragraph", {})
        run_spec = spec.get("run", {})
        constraint_spec = spec.get("constraints", {})

        for apply_target in requirement_rule.get("applies_to", []):
            kind = apply_target["kind"]
            if kind == "semantic_rule":
                target_rule = rules.get("semantic_rules", {}).get(apply_target["key"])
                if target_rule is None:
                    continue
                merge_non_none(target_rule.setdefault("paragraph", {}), paragraph_spec)
                merge_non_none(target_rule.setdefault("run", {}), run_spec)
            elif kind == "tracked_style":
                style_rule = find_style_rule(rules, apply_target["name"])
                if style_rule is None:
                    continue
                merge_non_none(style_rule.setdefault("paragraph", {}), paragraph_spec)
                merge_non_none(style_rule.setdefault("run", {}), run_spec)
            elif kind == "custom_paragraph_rule":
                custom_paragraph_rules.append(build_custom_paragraph_rule_v2(apply_target["name"], requirement_rule))
            elif kind == "custom_table_rule":
                custom_table_rules.append(build_custom_table_rule_v2(apply_target["name"], requirement_rule))
            elif kind == "header_rule":
                header_rule = rules.get("header_rules", {}).get(apply_target["name"])
                if header_rule is None:
                    continue
                merge_non_none(header_rule.setdefault("paragraph", {}), paragraph_spec)
                merge_non_none(header_rule.setdefault("run", {}), run_spec)
                merge_non_none(header_rule.setdefault("constraints", {}), constraint_spec)
                if header_rule.get("text"):
                    header_rule.setdefault("constraints", {}).setdefault("placeholder_texts", [header_rule["text"]])

    rules["custom_paragraph_rules"] = custom_paragraph_rules
    rules["custom_table_rules"] = custom_table_rules
    return rules


def summarize_text_requirement_rules(text_requirement_rules: list[dict[str, Any]]) -> dict[str, Any]:
    supported_rules = [rule for rule in text_requirement_rules if rule.get("supported")]
    manual_rules = [rule for rule in text_requirement_rules if not rule.get("supported")]
    return {
        "total": len(text_requirement_rules),
        "supported": len(supported_rules),
        "manual_only": len(manual_rules),
        "activated_targets": [rule.get("target") for rule in supported_rules if rule.get("target")],
        "manual_targets": [rule.get("target") for rule in manual_rules if rule.get("target")],
    }


def format_requirement_apply_target(apply_target: dict[str, Any]) -> str:
    kind = apply_target.get("kind")
    if kind == "semantic_rule":
        return f"semantic:{apply_target.get('key')}"
    if kind == "tracked_style":
        return f"style:{apply_target.get('name')}"
    if kind == "custom_paragraph_rule":
        return f"scope:{apply_target.get('name')}"
    if kind == "custom_table_rule":
        return f"table:{apply_target.get('name')}"
    if kind == "header_rule":
        return f"header:{apply_target.get('name')}"
    return kind or "unknown"


def render_text_requirements_markdown(
    text_requirement_rules: list[dict[str, Any]],
    summary: dict[str, Any],
) -> str:
    lines = [
        "# Text Requirement Rules",
        "",
        "These rules were parsed from template paragraphs that describe formatting requirements in prose.",
        "",
        f"- Total extracted notes: `{summary['total']}`",
        f"- Activated as machine-checkable rules: `{summary['supported']}`",
        f"- Stored for manual review only: `{summary['manual_only']}`",
        "",
        "| Source Paragraph | Target | Status | Applies To | Note |",
        "| --- | --- | --- | --- | --- |",
    ]
    for rule in text_requirement_rules:
        note = str(rule.get("text") or "").replace("|", "\\|")
        targets = ", ".join(format_requirement_apply_target(item) for item in rule.get("applies_to", [])) or "-"
        status = "active" if rule.get("supported") else "stored_only"
        lines.append(
            "| {index} | {target} | {status} | {targets} | {note} |".format(
                index=rule.get("source_paragraph_index"),
                target=rule.get("label") or rule.get("target") or "-",
                status=status,
                targets=targets,
                note=note[:160],
            )
        )
    lines.append("")
    return "\n".join(lines)


def register_rule_pack(root: Path, entry: dict[str, Any]) -> None:
    registry = load_registry(root)
    registry["rule_packs"] = [item for item in registry["rule_packs"] if item["path"] != entry["path"]]
    registry["rule_packs"].append(entry)
    if entry.get("validation_pass"):
        registry["latest_rule_pack"] = entry
    elif registry.get("latest_rule_pack", {}).get("path") == entry["path"]:
        valid_entries = [item for item in registry["rule_packs"] if item.get("validation_pass")]
        registry["latest_rule_pack"] = valid_entries[-1] if valid_entries else None
    save_registry(root, registry)


def ensure_validated_rule_pack(rule_pack_path: Path, pack: dict[str, Any]) -> None:
    validation = pack.get("manifest", {}).get("validation") or {}
    if validation.get("pass") and validation.get("comparison") == "docx-package-metadata":
        return
    report_path = validation.get("report_md") or validation.get("report_json")
    if validation.get("pass") and validation.get("comparison") != "docx-package-metadata":
        raise RuntimeError(
            f"Rule pack validation for {rule_pack_path} is stale. Rerun validate-rule-pack with DOCX package metadata diff enabled."
        )
    if report_path:
        raise RuntimeError(
            f"Rule pack validation failed for {rule_pack_path}. Inspect {report_path} and rerun validate-rule-pack."
        )
    raise RuntimeError(
        f"Rule pack validation has not passed for {rule_pack_path}. Run validate-rule-pack before audit or repair."
    )


def build_rule_pack(args: argparse.Namespace, root: Path) -> int:
    template_path = Path(args.template).resolve()
    if not template_path.exists():
        raise FileNotFoundError(f"Template DOCX not found: {template_path}")
    pack_name = args.name or slugify(template_path.stem)
    output_dir = (
        Path(args.output_dir).resolve()
        if args.output_dir
        else root / "artifacts" / "rule-packs" / pack_name
    )
    ensure_dir(output_dir)

    model = build_document_model(template_path)
    manifest = {
        "version": RULE_PACK_VERSION,
        "pipeline_version": PIPELINE_VERSION,
        "generated_at": utc_now(),
        "name": pack_name,
        "template": model["source"],
        "rule_pack_path": str(output_dir.resolve()),
    }
    rules = {
        "version": RULE_PACK_VERSION,
        "document_defaults": model["document_defaults"],
        "section_settings": model["sections"],
        "required_sections": model["required_sections"],
        "semantic_rules": model["semantic_rules"],
        "heading_rules": model["heading_rules"],
        "header_rules": header_rules_from_model(model),
        "cover_rules": {"blocks": model["cover_blocks"]},
        "tracked_styles": model["tracked_styles"],
        "requirement_notes": model["requirement_notes"],
        "template_blueprint": template_blueprint_from_model(model),
    }
    text_requirement_rules = build_text_requirement_rules_v2(model)
    rules = apply_text_requirement_rules_v2(rules, text_requirement_rules)
    text_requirement_summary = summarize_text_requirement_rules(text_requirement_rules)
    manifest["text_requirements"] = text_requirement_summary

    write_json(output_dir / "manifest.json", manifest)
    write_json(output_dir / "rules.json", rules)
    write_json(output_dir / "template-model.json", model)
    write_text(output_dir / "template-index.md", model["paragraph_index_markdown"])
    write_text(
        output_dir / "text-requirements.md",
        render_text_requirements_markdown(text_requirement_rules, text_requirement_summary),
    )
    package_manifest = extract_docx_package_snapshot(template_path, package_snapshot_dir(output_dir))
    write_json(package_manifest_path(output_dir), package_manifest)

    validation = validate_rule_pack_contents(
        root=root,
        rule_pack_path=output_dir,
        pack={"manifest": manifest, "rules": rules},
        template_path=template_path,
        template_model=model,
    )
    manifest["validation"] = {
        "validated_at": validation["generated_at"],
        "comparison": validation["comparison"],
        "pass": validation["summary"]["pass"],
        "errors": validation["summary"]["errors"],
        "warnings": validation["summary"]["warnings"],
        "rebuilt_template_path": validation["rebuilt_template_path"],
        "package_manifest": validation["package_manifest"],
        "report_json": validation["report_json"],
        "report_md": validation["report_md"],
    }
    write_json(output_dir / "manifest.json", manifest)

    entry = {
        "name": pack_name,
        "path": str(output_dir.resolve()),
        "template_path": str(template_path),
        "template_sha256": model["source"]["sha256"],
        "generated_at": manifest["generated_at"],
        "validation_comparison": validation["comparison"],
        "validation_pass": validation["summary"]["pass"],
        "validation_report_json": validation["report_json"],
        "validation_report_md": validation["report_md"],
    }
    register_rule_pack(root, entry)

    print(
        json.dumps(
            {
                "status": "ok" if validation["summary"]["pass"] else "invalid",
                "action": "build-rule-pack",
                "rule_pack_path": str(output_dir.resolve()),
                "template_path": str(template_path),
                "validation_report_json": validation["report_json"],
                "validation_report_md": validation["report_md"],
                "text_requirements": text_requirement_summary,
                "text_requirements_md": str((output_dir / "text-requirements.md").resolve()),
                "summary": validation["summary"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if validation["summary"]["pass"] else 1


def add_finding(
    findings: list[dict[str, Any]],
    severity: str,
    category: str,
    location: str,
    field: str,
    expected: Any,
    actual: Any,
    message: str,
) -> None:
    findings.append(
        {
            "severity": severity,
            "category": category,
            "location": location,
            "field": field,
            "expected": expected,
            "actual": actual,
            "message": message,
        }
    )


def values_differ(expected: Any, actual: Any, tolerance: float = 0.5) -> bool:
    if expected is None:
        return False
    if actual is None:
        return True
    if isinstance(expected, (int, float)) and isinstance(actual, (int, float)):
        return abs(float(expected) - float(actual)) > tolerance
    return str(expected).strip().lower() != str(actual).strip().lower()


def compare_mapping(
    findings: list[dict[str, Any]],
    *,
    expected: dict[str, Any],
    actual: dict[str, Any],
    category: str,
    location: str,
    severity: str,
    fields: list[str],
    tolerance_map: dict[str, float] | None = None,
) -> None:
    tolerance_map = tolerance_map or {}
    for field in fields:
        expected_value = expected.get(field)
        if expected_value is None:
            continue
        actual_value = actual.get(field)
        tolerance = tolerance_map.get(field, 0.5)
        if values_differ(expected_value, actual_value, tolerance=tolerance):
            add_finding(
                findings=findings,
                severity=severity,
                category=category,
                location=location,
                field=field,
                expected=expected_value,
                actual=actual_value,
                message=f"{field} does not match the rule pack.",
            )


def run_fields_for_record(record: dict[str, Any], expected_run: dict[str, Any]) -> list[str]:
    text = (record.get("text") or "").strip()
    has_cjk = bool(re.search(r"[\u4e00-\u9fff]", text))
    has_latin = bool(re.search(r"[A-Za-z0-9]", text))

    fields: list[str] = []
    if expected_run.get("size_pt") is not None:
        fields.append("size_pt")
    if expected_run.get("bold") is not None:
        fields.append("bold")
    if expected_run.get("italic") is not None:
        fields.append("italic")
    if has_cjk and expected_run.get("east_asia_font") is not None:
        fields.append("east_asia_font")
    if has_latin and not has_cjk and expected_run.get("font_name") is not None:
        fields.append("font_name")
    if not has_cjk and not has_latin and expected_run.get("font_name") is not None:
        fields.append("font_name")

    ordered: list[str] = []
    for field in fields:
        if field not in ordered:
            ordered.append(field)
    return ordered


def model_by_semantic_label(model: dict[str, Any]) -> dict[str, dict[str, Any]]:
    mapping: dict[str, dict[str, Any]] = {}
    for paragraph in model["paragraphs"]:
        label = paragraph.get("semantic_label")
        if label and label not in mapping:
            mapping[label] = paragraph
    return mapping


def paragraphs_in_scope(model: dict[str, Any], scope: dict[str, Any]) -> list[dict[str, Any]]:
    paragraphs = model.get("paragraphs", [])
    start_label = scope.get("start_label")
    end_labels = set(scope.get("end_labels", []))

    if start_label:
        start_index = next(
            (paragraph["index"] for paragraph in paragraphs if paragraph.get("semantic_label") == start_label),
            None,
        )
        if start_index is None:
            return []
    else:
        start_index = -1

    end_index = len(paragraphs)
    if end_labels:
        for paragraph in paragraphs:
            if paragraph["index"] <= start_index:
                continue
            if paragraph.get("semantic_label") in end_labels:
                end_index = paragraph["index"]
                break
    return paragraphs[start_index + 1 : end_index]


def is_code_like_paragraph(paragraph: dict[str, Any], scope: dict[str, Any]) -> bool:
    text = (paragraph.get("text") or "").strip()
    if not text:
        return False
    max_text_length = scope.get("max_text_length")
    if max_text_length is not None and len(text) > int(max_text_length):
        return False

    run_fonts = {
        (run.get("run") or {}).get("font_name")
        for run in paragraph.get("runs", [])
        if (run.get("text") or "").strip()
    }
    expected_fonts = set(scope.get("run_font_names", []))
    if expected_fonts.intersection(font for font in run_fonts if font):
        return True

    code_markers = [
        "{",
        "}",
        ";",
        "()",
        "public ",
        "private ",
        "class ",
        "return ",
        " if ",
        "if(",
        "for(",
        "while(",
        "switch(",
        "case ",
    ]
    marker_hits = sum(1 for marker in code_markers if marker in text)
    ascii_chars = sum(1 for ch in text if ord(ch) < 128 and not ch.isspace())
    ascii_ratio = ascii_chars / max(len(text), 1)
    return marker_hits >= 2 or (marker_hits >= 1 and ascii_ratio >= 0.45)


def paragraph_matches_scope(paragraph: dict[str, Any], scope: dict[str, Any]) -> bool:
    if not scope.get("include_empty") and paragraph.get("is_empty"):
        return False
    if scope.get("exclude_any_semantic_label") and paragraph.get("semantic_label"):
        return False
    allowed_style_names = set(scope.get("allowed_style_names", []))
    if allowed_style_names and paragraph.get("style_name") not in allowed_style_names:
        return False
    pattern = scope.get("text_regex")
    if pattern and not re.search(pattern, paragraph.get("text") or ""):
        return False
    if scope.get("heuristic") == "code_like" and not is_code_like_paragraph(paragraph, scope):
        return False
    return True


def paragraphs_for_custom_rule(model: dict[str, Any], rule: dict[str, Any]) -> list[dict[str, Any]]:
    scope = rule.get("scope", {})
    return [
        paragraph
        for paragraph in paragraphs_in_scope(model, scope)
        if paragraph_matches_scope(paragraph, scope)
    ]


def table_paragraphs_for_custom_rule(model: dict[str, Any], rule: dict[str, Any]) -> list[dict[str, Any]]:
    scope = rule.get("scope", {})
    return [
        paragraph
        for paragraph in model.get("table_paragraphs", [])
        if paragraph_matches_scope(paragraph, scope)
    ]


def header_paragraphs_for_rule(model: dict[str, Any], rule: dict[str, Any]) -> list[dict[str, Any]]:
    header_kind = rule.get("header_kind", "default")
    return [
        paragraph
        for paragraph in model.get("headers", [])
        if paragraph.get("header_kind") == header_kind and not paragraph.get("is_empty")
    ]


def add_requirement_constraint_findings(
    findings: list[dict[str, Any]],
    *,
    rule: dict[str, Any],
    actual: dict[str, Any],
    location: str,
) -> None:
    constraints = rule.get("constraints", {})
    text = (actual.get("text") or "").strip()
    text_regex = constraints.get("text_regex")
    if text_regex and text and not re.search(text_regex, text):
        add_finding(
            findings,
            severity="warning",
            category="requirement-content",
            location=location,
            field="text",
            expected=text_regex,
            actual=text,
            message="Text does not satisfy the requirement pattern.",
        )

    placeholder_texts = constraints.get("placeholder_texts", [])
    if constraints.get("must_not_equal_expected_text") and text in placeholder_texts:
        add_finding(
            findings,
            severity="warning",
            category="requirement-content",
            location=location,
            field="text",
            expected="customized header text",
            actual=text,
            message="Text still matches the template placeholder and likely was not customized.",
        )


def audit_against_rule_pack(
    model: dict[str, Any],
    rule_pack: dict[str, Any],
) -> dict[str, Any]:
    rules = rule_pack["rules"]
    findings: list[dict[str, Any]] = []
    actual_semantics = model_by_semantic_label(model)
    actual_sections = model.get("required_sections", [])
    expected_sections = rules.get("required_sections", [])

    expected_section = (rules.get("section_settings") or [{}])[0]
    actual_section = (model.get("sections") or [{}])[0]
    compare_mapping(
        findings,
        expected=expected_section,
        actual=actual_section,
        category="section",
        location="document section 0",
        severity="error",
        fields=[
            "orientation",
            "page_width_pt",
            "page_height_pt",
            "left_margin_pt",
            "right_margin_pt",
            "top_margin_pt",
            "bottom_margin_pt",
            "header_distance_pt",
            "footer_distance_pt",
        ],
        tolerance_map={field: 1.0 for field in [
            "page_width_pt",
            "page_height_pt",
            "left_margin_pt",
            "right_margin_pt",
            "top_margin_pt",
            "bottom_margin_pt",
            "header_distance_pt",
            "footer_distance_pt",
        ]},
    )

    for label in expected_sections:
        if label not in actual_sections:
            add_finding(
                findings,
                severity="error",
                category="structure",
                location="document",
                field="required_sections",
                expected=label,
                actual=None,
                message=f"Required section '{label}' is missing.",
            )
    existing_indices = [actual_sections.index(label) for label in expected_sections if label in actual_sections]
    if existing_indices != sorted(existing_indices):
        add_finding(
            findings,
            severity="error",
            category="structure",
            location="document",
            field="section_order",
            expected=expected_sections,
            actual=actual_sections,
            message="Detected thesis sections are out of order.",
        )

    for label, expected_rule in rules.get("semantic_rules", {}).items():
        actual = actual_semantics.get(label)
        if actual is None:
            continue
        location = f"paragraph {actual['index']} ({label})"
        compare_mapping(
            findings,
            expected={"style_name": expected_rule.get("style_name")},
            actual={"style_name": actual.get("style_name")},
            category="semantic-style",
            location=location,
            severity="error",
            fields=["style_name"],
        )
        compare_mapping(
            findings,
            expected=expected_rule.get("paragraph", {}),
            actual=actual.get("paragraph", {}),
            category="semantic-paragraph",
            location=location,
            severity="warning",
            fields=[
                "alignment",
                "space_before_pt",
                "space_after_pt",
                "line_spacing_pt",
                "first_line_indent_pt",
            ],
            tolerance_map={
                "space_before_pt": 1.0,
                "space_after_pt": 1.0,
                "line_spacing_pt": 1.0,
                "first_line_indent_pt": 1.0,
            },
        )
        compare_mapping(
            findings,
            expected=expected_rule.get("run", {}),
            actual=actual.get("run", {}),
            category="semantic-run",
            location=location,
            severity="warning",
            fields=run_fields_for_record(actual, expected_rule.get("run", {})),
            tolerance_map={"size_pt": 0.5},
        )

    for header_name, expected_rule in rules.get("header_rules", {}).items():
        matched_headers = header_paragraphs_for_rule(model, expected_rule)
        if not matched_headers:
            add_finding(
                findings,
                severity="warning",
                category="requirement-header-paragraph",
                location=f"header kind {header_name}",
                field="exists",
                expected="present",
                actual="missing",
                message="Required header paragraph is missing.",
            )
            continue
        for header in matched_headers:
            location = (
                f"header section {header['section_index']} {header['header_kind']} "
                f"paragraph {header['header_paragraph_index']}"
            )
            compare_mapping(
                findings,
                expected=expected_rule.get("paragraph", {}),
                actual=header.get("paragraph", {}),
                category="requirement-header-paragraph",
                location=location,
                severity="warning",
                fields=[
                    "alignment",
                    "space_before_pt",
                    "space_after_pt",
                    "line_spacing_pt",
                    "first_line_indent_pt",
                ],
                tolerance_map={
                    "space_before_pt": 1.0,
                    "space_after_pt": 1.0,
                    "line_spacing_pt": 1.0,
                    "first_line_indent_pt": 1.0,
                },
            )
            compare_mapping(
                findings,
                expected=expected_rule.get("run", {}),
                actual=header.get("run", {}),
                category="requirement-header-run",
                location=location,
                severity="warning",
                fields=run_fields_for_record(header, expected_rule.get("run", {})),
                tolerance_map={"size_pt": 0.5},
            )
            add_requirement_constraint_findings(
                findings,
                rule=expected_rule,
                actual=header,
                location=location,
            )

    for custom_rule in rules.get("custom_paragraph_rules", []):
        matched_paragraphs = paragraphs_for_custom_rule(model, custom_rule)
        for paragraph in matched_paragraphs:
            location = f"paragraph {paragraph['index']} ({custom_rule['name']})"
            compare_mapping(
                findings,
                expected=custom_rule.get("paragraph", {}),
                actual=paragraph.get("paragraph", {}),
                category="requirement-paragraph",
                location=location,
                severity="warning",
                fields=[
                    "alignment",
                    "space_before_pt",
                    "space_after_pt",
                    "line_spacing_pt",
                    "first_line_indent_pt",
                ],
                tolerance_map={
                    "space_before_pt": 1.0,
                    "space_after_pt": 1.0,
                    "line_spacing_pt": 1.0,
                    "first_line_indent_pt": 1.0,
                },
            )
            compare_mapping(
                findings,
                expected=custom_rule.get("run", {}),
                actual=paragraph.get("run", {}),
                category="requirement-run",
                location=location,
                severity="warning",
                fields=run_fields_for_record(paragraph, custom_rule.get("run", {})),
                tolerance_map={"size_pt": 0.5},
            )
            add_requirement_constraint_findings(
                findings,
                rule=custom_rule,
                actual=paragraph,
                location=location,
            )

    for table_rule in rules.get("custom_table_rules", []):
        matched_paragraphs = table_paragraphs_for_custom_rule(model, table_rule)
        for paragraph in matched_paragraphs:
            location = (
                f"table {paragraph['table_index']} row {paragraph['row_index']} "
                f"col {paragraph['column_index']} paragraph {paragraph['cell_paragraph_index']}"
            )
            compare_mapping(
                findings,
                expected=table_rule.get("paragraph", {}),
                actual=paragraph.get("paragraph", {}),
                category="requirement-table-paragraph",
                location=location,
                severity="warning",
                fields=[
                    "alignment",
                    "space_before_pt",
                    "space_after_pt",
                    "line_spacing_pt",
                    "first_line_indent_pt",
                ],
                tolerance_map={
                    "space_before_pt": 1.0,
                    "space_after_pt": 1.0,
                    "line_spacing_pt": 1.0,
                    "first_line_indent_pt": 1.0,
                },
            )
            compare_mapping(
                findings,
                expected=table_rule.get("run", {}),
                actual=paragraph.get("run", {}),
                category="requirement-table-run",
                location=location,
                severity="warning",
                fields=run_fields_for_record(paragraph, table_rule.get("run", {})),
                tolerance_map={"size_pt": 0.5},
            )
            add_requirement_constraint_findings(
                findings,
                rule=table_rule,
                actual=paragraph,
                location=location,
            )

    expected_cover = rules.get("cover_rules", {}).get("blocks", [])
    actual_cover = model.get("cover_blocks", [])
    if len(actual_cover) < len(expected_cover):
        add_finding(
            findings,
            severity="error",
            category="cover",
            location="cover",
            field="block_count",
            expected=len(expected_cover),
            actual=len(actual_cover),
            message="Cover/front-matter block count is smaller than the template.",
        )
    for index, expected_block in enumerate(expected_cover[: len(actual_cover)]):
        actual_block = actual_cover[index]
        location = f"cover block {index} (paragraph {actual_block['index']})"
        compare_mapping(
            findings,
            expected={
                "style_name": expected_block.get("style_name"),
                "leading_blank_paragraphs": expected_block.get("leading_blank_paragraphs"),
            },
            actual={
                "style_name": actual_block.get("style_name"),
                "leading_blank_paragraphs": actual_block.get("leading_blank_paragraphs"),
            },
            category="cover",
            location=location,
            severity="warning",
            fields=["style_name", "leading_blank_paragraphs"],
        )
        compare_mapping(
            findings,
            expected=expected_block.get("paragraph", {}),
            actual=actual_block.get("paragraph", {}),
            category="cover-paragraph",
            location=location,
            severity="warning",
            fields=["alignment", "space_before_pt", "space_after_pt"],
            tolerance_map={"space_before_pt": 1.0, "space_after_pt": 1.0},
        )
        compare_mapping(
            findings,
            expected=expected_block.get("run", {}),
            actual=actual_block.get("run", {}),
            category="cover-run",
            location=location,
            severity="warning",
            fields=run_fields_for_record(actual_block, expected_block.get("run", {})),
            tolerance_map={"size_pt": 0.5},
        )

    actual_styles = {style["style_id"]: style for style in model.get("styles", {}).values()}
    actual_style_names = {style["name"]: style for style in model.get("styles", {}).values() if style.get("name")}
    for expected_style in rules.get("tracked_styles", []):
        actual_style = actual_styles.get(expected_style.get("style_id")) or actual_style_names.get(
            expected_style.get("name")
        )
        if actual_style is None:
            add_finding(
                findings,
                severity="warning",
                category="style-definition",
                location=expected_style.get("name") or expected_style.get("style_id") or "style",
                field="exists",
                expected="present",
                actual="missing",
                message="Tracked style from template is missing in the thesis document.",
            )
            continue
        location = expected_style.get("name") or expected_style.get("style_id") or "style"
        compare_mapping(
            findings,
            expected=expected_style.get("paragraph", {}),
            actual=actual_style.get("paragraph", {}),
            category="style-definition",
            location=location,
            severity="warning",
            fields=["alignment", "space_before_pt", "space_after_pt", "line_spacing_pt"],
            tolerance_map={"space_before_pt": 1.0, "space_after_pt": 1.0, "line_spacing_pt": 1.0},
        )
        compare_mapping(
            findings,
            expected=expected_style.get("run", {}),
            actual=actual_style.get("run", {}),
            category="style-definition",
            location=location,
            severity="warning",
            fields=["font_name", "east_asia_font", "size_pt", "bold", "italic"],
            tolerance_map={"size_pt": 0.5},
        )

    summary = {
        "errors": sum(1 for item in findings if item["severity"] == "error"),
        "warnings": sum(1 for item in findings if item["severity"] == "warning"),
        "info": sum(1 for item in findings if item["severity"] == "info"),
    }
    summary["pass"] = summary["errors"] == 0
    return {"summary": summary, "findings": findings}


SEVERITY_LABELS = {
    "error": "错误",
    "warning": "警告",
    "info": "提示",
}

CATEGORY_LABELS = {
    "section": "页面设置",
    "structure": "结构",
    "semantic-style": "语义样式",
    "semantic-paragraph": "语义段落",
    "semantic-run": "语义字体",
    "requirement-paragraph": "文字要求段落",
    "requirement-run": "文字要求字体",
    "cover": "封面",
    "cover-paragraph": "封面段落",
    "cover-run": "封面字体",
    "style-definition": "样式定义",
}

CATEGORY_LABELS.update(
    {
        "requirement-header-paragraph": "页眉要求段落",
        "requirement-header-run": "页眉要求字体",
        "requirement-table-paragraph": "表格要求段落",
        "requirement-table-run": "表格要求字体",
        "requirement-content": "文本要求",
    }
)

FIELD_LABELS = {
    "orientation": "页面方向",
    "page_width_pt": "页面宽度(pt)",
    "page_height_pt": "页面高度(pt)",
    "left_margin_pt": "左页边距(pt)",
    "right_margin_pt": "右页边距(pt)",
    "top_margin_pt": "上页边距(pt)",
    "bottom_margin_pt": "下页边距(pt)",
    "header_distance_pt": "页眉距离(pt)",
    "footer_distance_pt": "页脚距离(pt)",
    "required_sections": "必需章节",
    "section_order": "章节顺序",
    "style_name": "样式名",
    "alignment": "对齐方式",
    "space_before_pt": "段前距(pt)",
    "space_after_pt": "段后距(pt)",
    "line_spacing_pt": "行距(pt)",
    "first_line_indent_pt": "首行缩进(pt)",
    "font_name": "西文字体",
    "east_asia_font": "中文字体",
    "size_pt": "字号(pt)",
    "bold": "加粗",
    "italic": "斜体",
    "block_count": "封面块数量",
    "leading_blank_paragraphs": "前导空段数",
    "exists": "存在性",
}

FIELD_LABELS.update({"text": "文本"})

SEMANTIC_LABELS = {
    "abstract_cn": "中文摘要",
    "abstract_en": "英文摘要",
    "keywords_cn": "中文关键词",
    "keywords_en": "英文关键词",
    "toc": "目录",
    "references": "参考文献",
    "acknowledgements": "致谢",
    "appendix": "附录",
    "chapter_heading": "章标题",
    "section_heading": "节标题",
    "abstract_cn_body": "中文摘要正文",
    "abstract_en_body": "英文摘要正文",
    "body_paragraph": "正文段落",
    "reference_entry": "参考文献条目",
    "heading_level_1": "一级标题要求",
    "heading_level_2": "二级标题要求",
    "heading_level_3": "三级标题要求",
}


def translate_semantic_label(value: str) -> str:
    return SEMANTIC_LABELS.get(value, value)


def translate_severity(value: str) -> str:
    return SEVERITY_LABELS.get(value, value)


def translate_category(value: str) -> str:
    return CATEGORY_LABELS.get(value, value)


def translate_field(value: str) -> str:
    return FIELD_LABELS.get(value, value)


def format_report_value(value: Any) -> str:
    if value is None:
        return "缺失"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, list):
        return "、".join(format_report_value(item) for item in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    if isinstance(value, str):
        return translate_semantic_label(value)
    return str(value)


def translate_location(value: str) -> str:
    paragraph_match = re.fullmatch(r"paragraph (\d+) \(([^)]+)\)", value)
    if paragraph_match:
        index, label = paragraph_match.groups()
        return f"段落 {index}（{translate_semantic_label(label)}）"

    paragraph_run_match = re.fullmatch(r"paragraph (\d+) / run (\d+)", value)
    if paragraph_run_match:
        paragraph_index, run_index = paragraph_run_match.groups()
        return f"段落 {paragraph_index} / run {run_index}"

    paragraph_only_match = re.fullmatch(r"paragraph (\d+)", value)
    if paragraph_only_match:
        return f"段落 {paragraph_only_match.group(1)}"

    cover_match = re.fullmatch(r"cover block (\d+) \(paragraph (\d+)\)", value)
    if cover_match:
        block_index, paragraph_index = cover_match.groups()
        return f"封面块 {block_index}（段落 {paragraph_index}）"

    if value == "document section 0":
        return "文档节 0"
    if value == "document":
        return "整篇文档"
    return value


def translate_message(finding: dict[str, Any]) -> str:
    field_label = translate_field(finding["field"])
    message = finding["message"]

    if message == f"{finding['field']} does not match the rule pack.":
        return f"{field_label}与规则包要求不一致。"

    required_match = re.fullmatch(r"Required section '([^']+)' is missing\.", message)
    if required_match:
        return f"缺少必需章节“{translate_semantic_label(required_match.group(1))}”。"

    if message == "Detected thesis sections are out of order.":
        return "检测到论文主要章节顺序与规则包不一致。"

    if message == "Cover/front-matter block count is smaller than the template.":
        return "封面或前置部分的块数量少于模板。"

    if message == "Tracked style from template is missing in the thesis document.":
        return "论文中缺少模板要求跟踪的样式定义。"

    return message


def translate_location_v2(value: str) -> str:
    translated = translate_location(value)
    if translated != value:
        return translated

    header_match = re.fullmatch(r"header section (\d+) ([a-z_]+) paragraph (\d+)", value)
    if header_match:
        section_index, header_kind, paragraph_index = header_match.groups()
        return f"页眉 section {section_index} / {header_kind} / 段落 {paragraph_index}"

    header_kind_match = re.fullmatch(r"header kind ([a-z_]+)", value)
    if header_kind_match:
        return f"页眉类型 {header_kind_match.group(1)}"

    table_match = re.fullmatch(r"table (\d+) row (\d+) col (\d+) paragraph (\d+)", value)
    if table_match:
        table_index, row_index, column_index, paragraph_index = table_match.groups()
        return f"表格 {table_index} / 行 {row_index} / 列 {column_index} / 段落 {paragraph_index}"

    return value


def translate_message_v2(finding: dict[str, Any]) -> str:
    translated = translate_message(finding)
    if translated != finding["message"]:
        return translated
    if finding["message"] == "Required header paragraph is missing.":
        return "缺少规则包要求的页眉段落。"
    if finding["message"] == "Text does not satisfy the requirement pattern.":
        return "文本未满足规则包要求的编号或文本模式。"
    if finding["message"] == "Text still matches the template placeholder and likely was not customized.":
        return "文本仍然是模板占位内容，疑似未按论文信息替换。"
    return finding["message"]


def snippet_text(text: str, width: int = 18) -> dict[str, str]:
    text = (text or "").strip()
    if not text:
        text = "<EMPTY>"
    return {
        "prefix": text[:width],
        "suffix": text[-width:] if len(text) > width else text,
    }


def paragraph_context_from_model(paragraph: dict[str, Any]) -> dict[str, Any]:
    snippet = snippet_text(paragraph.get("text") or "")
    return {
        "paragraph_index": paragraph["index"],
        "style_name": paragraph.get("style_name"),
        "text_prefix": snippet["prefix"],
        "text_suffix": snippet["suffix"],
        "text_length": len((paragraph.get("text") or "").strip()),
    }


def build_style_examples(model: dict[str, Any], style_name: str, limit: int = 5) -> dict[str, Any]:
    matched = [
        paragraph
        for paragraph in model.get("paragraphs", [])
        if paragraph.get("style_name") == style_name
    ]
    examples = []
    for paragraph in matched[:limit]:
        context = paragraph_context_from_model(paragraph)
        examples.append(
            {
                "paragraph_index": context["paragraph_index"],
                "text_prefix": context["text_prefix"],
                "text_suffix": context["text_suffix"],
            }
        )
    return {
        "style_name": style_name,
        "matched_paragraph_count": len(matched),
        "examples": examples,
    }


def enrich_finding_context(finding: dict[str, Any], model: dict[str, Any]) -> dict[str, Any]:
    location = finding["location"]

    paragraph_match = re.fullmatch(r"paragraph (\d+) \(([^)]+)\)", location)
    if paragraph_match:
        paragraph_index = int(paragraph_match.group(1))
        paragraphs = model.get("paragraphs", [])
        if 0 <= paragraph_index < len(paragraphs):
            context = paragraph_context_from_model(paragraphs[paragraph_index])
            context["semantic_label"] = paragraph_match.group(2)
            return context
        return {}

    paragraph_run_match = re.fullmatch(r"paragraph (\d+) / run (\d+)", location)
    if paragraph_run_match:
        paragraph_index = int(paragraph_run_match.group(1))
        paragraphs = model.get("paragraphs", [])
        if 0 <= paragraph_index < len(paragraphs):
            context = paragraph_context_from_model(paragraphs[paragraph_index])
            context["run_index"] = int(paragraph_run_match.group(2))
            return context
        return {}

    paragraph_only_match = re.fullmatch(r"paragraph (\d+)", location)
    if paragraph_only_match:
        paragraph_index = int(paragraph_only_match.group(1))
        paragraphs = model.get("paragraphs", [])
        if 0 <= paragraph_index < len(paragraphs):
            return paragraph_context_from_model(paragraphs[paragraph_index])
        return {}

    cover_match = re.fullmatch(r"cover block (\d+) \(paragraph (\d+)\)", location)
    if cover_match:
        paragraph_index = int(cover_match.group(2))
        paragraphs = model.get("paragraphs", [])
        if 0 <= paragraph_index < len(paragraphs):
            context = paragraph_context_from_model(paragraphs[paragraph_index])
            context["cover_block_index"] = int(cover_match.group(1))
            return context
        return {}

    if finding["category"] == "style-definition":
        return build_style_examples(model, location)

    return {}


def enrich_findings_with_context(findings: list[dict[str, Any]], model: dict[str, Any]) -> list[dict[str, Any]]:
    enriched = []
    for finding in findings:
        item = dict(finding)
        context = enrich_finding_context(item, model)
        if context:
            item["context"] = context
        enriched.append(item)
    return enriched


def header_context_from_model(paragraph: dict[str, Any]) -> dict[str, Any]:
    context = paragraph_context_from_model(paragraph)
    context["header_section_index"] = paragraph.get("section_index")
    context["header_kind"] = paragraph.get("header_kind")
    context["header_paragraph_index"] = paragraph.get("header_paragraph_index")
    return context


def table_context_from_model(paragraph: dict[str, Any]) -> dict[str, Any]:
    context = paragraph_context_from_model(paragraph)
    context["table_index"] = paragraph.get("table_index")
    context["row_index"] = paragraph.get("row_index")
    context["column_index"] = paragraph.get("column_index")
    context["cell_paragraph_index"] = paragraph.get("cell_paragraph_index")
    return context


def enrich_finding_context_v2(finding: dict[str, Any], model: dict[str, Any]) -> dict[str, Any]:
    context = enrich_finding_context(finding, model)
    if context:
        return context

    location = finding["location"]
    header_match = re.fullmatch(r"header section (\d+) ([a-z_]+) paragraph (\d+)", location)
    if header_match:
        section_index = int(header_match.group(1))
        header_kind = header_match.group(2)
        paragraph_index = int(header_match.group(3))
        for paragraph in model.get("headers", []):
            if (
                paragraph.get("section_index") == section_index
                and paragraph.get("header_kind") == header_kind
                and paragraph.get("header_paragraph_index") == paragraph_index
            ):
                return header_context_from_model(paragraph)
        return {}

    table_match = re.fullmatch(r"table (\d+) row (\d+) col (\d+) paragraph (\d+)", location)
    if table_match:
        table_index = int(table_match.group(1))
        row_index = int(table_match.group(2))
        column_index = int(table_match.group(3))
        paragraph_index = int(table_match.group(4))
        for paragraph in model.get("table_paragraphs", []):
            if (
                paragraph.get("table_index") == table_index
                and paragraph.get("row_index") == row_index
                and paragraph.get("column_index") == column_index
                and paragraph.get("cell_paragraph_index") == paragraph_index
            ):
                return table_context_from_model(paragraph)
        return {}

    return {}


def enrich_findings_with_context_v2(findings: list[dict[str, Any]], model: dict[str, Any]) -> list[dict[str, Any]]:
    enriched = []
    for finding in findings:
        item = dict(finding)
        context = enrich_finding_context_v2(item, model)
        if context:
            item["context"] = context
        enriched.append(item)
    return enriched


def render_audit_markdown(report: dict[str, Any], root: Path) -> str:
    lines = [
        "# 论文格式审查报告",
        "",
        f"- 生成时间：`{report['generated_at']}`",
        f"- 规则包：`{relative_path(Path(report['rule_pack_path']), root)}`",
        f"- 论文文件：`{relative_path(Path(report['document_path']), root)}`",
        f"- 是否通过：`{'是' if report['summary']['pass'] else '否'}`",
        f"- 错误数：`{report['summary']['errors']}`",
        f"- 警告数：`{report['summary']['warnings']}`",
        "",
        "## 问题明细",
        "",
    ]
    if not report["findings"]:
        lines.extend(["未发现问题。", ""])
        return "\n".join(lines)
    for finding in report["findings"]:
        lines.append(
            "- [{severity}] {category} | {location} | {field} | 期望值 `{expected}` | 实际值 `{actual}` | {message}".format(
                severity=translate_severity(finding["severity"]),
                category=translate_category(finding["category"]),
                location=translate_location(finding["location"]),
                field=translate_field(finding["field"]),
                expected=format_report_value(finding["expected"]),
                actual=format_report_value(finding["actual"]),
                message=translate_message(finding),
            )
        )
        context = finding.get("context") or {}
        paragraph_index = context.get("paragraph_index")
        if paragraph_index is not None:
            style_name = context.get("style_name") or "未标注"
            lines.append(f"  段落号：`{paragraph_index}`，当前样式：`{style_name}`")
            lines.append(
                "  段落摘录：前 18 字 `{prefix}` | 后 18 字 `{suffix}`".format(
                    prefix=context.get("text_prefix", ""),
                    suffix=context.get("text_suffix", ""),
                )
            )
        elif finding["category"] == "style-definition":
            lines.append(
                "  影响范围：样式 `{style_name}`，命中段落数 `{count}`".format(
                    style_name=context.get("style_name", finding["location"]),
                    count=context.get("matched_paragraph_count", 0),
                )
            )
            examples = context.get("examples") or []
            if examples:
                example_text = "；".join(
                    "段落 {paragraph_index}：前 18 字 `{text_prefix}` | 后 18 字 `{text_suffix}`".format(
                        paragraph_index=example["paragraph_index"],
                        text_prefix=example["text_prefix"],
                        text_suffix=example["text_suffix"],
                    )
                    for example in examples
                )
                lines.append(f"  示例段落：{example_text}")
    lines.append("")
    return "\n".join(lines)


def render_audit_markdown_v2(report: dict[str, Any], root: Path) -> str:
    lines = [
        "# 论文格式审查报告",
        "",
        f"- 生成时间：`{report['generated_at']}`",
        f"- 规则包：`{relative_path(Path(report['rule_pack_path']), root)}`",
        f"- 论文文件：`{relative_path(Path(report['document_path']), root)}`",
        f"- 是否通过：`{'是' if report['summary']['pass'] else '否'}`",
        f"- 错误数：`{report['summary']['errors']}`",
        f"- 警告数：`{report['summary']['warnings']}`",
        "",
        "## 问题明细",
        "",
    ]
    if not report["findings"]:
        lines.extend(["未发现问题。", ""])
        return "\n".join(lines)

    for finding in report["findings"]:
        lines.append(
            "- [{severity}] {category} | {location} | {field} | 期望值 `{expected}` | 实际值 `{actual}` | {message}".format(
                severity=translate_severity(finding["severity"]),
                category=translate_category(finding["category"]),
                location=translate_location_v2(finding["location"]),
                field=translate_field(finding["field"]),
                expected=format_report_value(finding["expected"]),
                actual=format_report_value(finding["actual"]),
                message=translate_message_v2(finding),
            )
        )
        context = finding.get("context") or {}
        paragraph_index = context.get("paragraph_index")
        if paragraph_index is not None:
            lines.append(
                "  段落位置：`{index}`，当前样式：`{style}`".format(
                    index=paragraph_index,
                    style=context.get("style_name") or "未标注",
                )
            )
            lines.append(
                "  文本范围：前 18 字 `{prefix}` | 后 18 字 `{suffix}`".format(
                    prefix=context.get("text_prefix", ""),
                    suffix=context.get("text_suffix", ""),
                )
            )
            continue
        if context.get("header_section_index") is not None:
            lines.append(
                "  页眉位置：section `{section}` / `{kind}` / 段落 `{paragraph}`".format(
                    section=context.get("header_section_index"),
                    kind=context.get("header_kind"),
                    paragraph=context.get("header_paragraph_index"),
                )
            )
            lines.append(
                "  文本范围：前 18 字 `{prefix}` | 后 18 字 `{suffix}`".format(
                    prefix=context.get("text_prefix", ""),
                    suffix=context.get("text_suffix", ""),
                )
            )
            continue
        if context.get("table_index") is not None:
            lines.append(
                "  表格位置：表 `{table}` / 行 `{row}` / 列 `{column}` / 段落 `{paragraph}`".format(
                    table=context.get("table_index"),
                    row=context.get("row_index"),
                    column=context.get("column_index"),
                    paragraph=context.get("cell_paragraph_index"),
                )
            )
            lines.append(
                "  文本范围：前 18 字 `{prefix}` | 后 18 字 `{suffix}`".format(
                    prefix=context.get("text_prefix", ""),
                    suffix=context.get("text_suffix", ""),
                )
            )
            continue
        if finding["category"] == "style-definition":
            lines.append(
                "  影响范围：样式 `{style}`，命中段落数 `{count}`".format(
                    style=context.get("style_name", finding["location"]),
                    count=context.get("matched_paragraph_count", 0),
                )
            )
            examples = context.get("examples") or []
            if examples:
                example_text = "；".join(
                    "段落 {paragraph_index}：前 18 字 `{text_prefix}` | 后 18 字 `{text_suffix}`".format(
                        paragraph_index=example["paragraph_index"],
                        text_prefix=example["text_prefix"],
                        text_suffix=example["text_suffix"],
                    )
                    for example in examples
                )
                lines.append(f"  示例段落：{example_text}")
    lines.append("")
    return "\n".join(lines)


def audit_document(args: argparse.Namespace, root: Path) -> int:
    document_path = Path(args.document).resolve()
    if not document_path.exists():
        raise FileNotFoundError(f"Thesis DOCX not found: {document_path}")
    rule_pack_path = find_rule_pack_path(root, args.rule_pack)
    if rule_pack_path is None:
        raise FileNotFoundError("No rule pack found. Build one first or pass --rule-pack.")
    pack = load_rule_pack(rule_pack_path)
    ensure_validated_rule_pack(rule_pack_path, pack)
    model = build_document_model(document_path)
    audit = audit_against_rule_pack(model, pack)

    report_dir = (
        Path(args.report_dir).resolve()
        if args.report_dir
        else root / "artifacts" / "audits" / path_slug(document_path, root)
    )
    ensure_dir(report_dir)
    report = {
        "version": RULE_PACK_VERSION,
        "generated_at": utc_now(),
        "pipeline_version": PIPELINE_VERSION,
        "rule_pack_path": str(rule_pack_path.resolve()),
        "document_path": str(document_path.resolve()),
        "document_sha256": model["source"]["sha256"],
        "summary": audit["summary"],
        "findings": enrich_findings_with_context_v2(audit["findings"], model),
    }
    json_path = report_dir / "audit-report.json"
    md_path = report_dir / "audit-report.md"
    write_json(json_path, report)
    write_text(md_path, render_audit_markdown_v2(report, root))

    registry = load_registry(root)
    entry = {
        "generated_at": report["generated_at"],
        "document_path": str(document_path.resolve()),
        "rule_pack_path": str(rule_pack_path.resolve()),
        "report_json": str(json_path.resolve()),
        "report_md": str(md_path.resolve()),
        "summary": report["summary"],
    }
    registry["audits"].append(entry)
    registry["latest_audit"] = entry
    save_registry(root, registry)

    print(
        json.dumps(
            {
                "status": "ok",
                "action": "audit",
                "report_json": str(json_path.resolve()),
                "report_md": str(md_path.resolve()),
                "summary": report["summary"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


def pt_to_twips(value: float | int | None) -> str | None:
    if value is None:
        return None
    return str(int(round(float(value) * 20.0)))


def pt_to_half_points(value: float | int | None) -> str | None:
    if value is None:
        return None
    return str(int(round(float(value) * 2.0)))


def ensure_child(element: Any, tag: str) -> Any:
    child = element.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        element.append(child)
    return child


def set_on_off(parent: Any, tag: str, value: bool | None) -> None:
    if value is None:
        return
    child = ensure_child(parent, tag)
    child.set(qn("w:val"), "1" if value else "0")


def set_decimal_value(parent: Any, tag: str, value: str | None) -> None:
    if value is None:
        return
    child = ensure_child(parent, tag)
    child.set(qn("w:val"), value)


def apply_run_spec(font_proxy: Any, xml_parent: Any, spec: dict[str, Any]) -> None:
    if spec.get("font_name") is not None:
        font_proxy.name = spec["font_name"]
    if spec.get("size_pt") is not None:
        font_proxy.size = Pt(spec["size_pt"])
    if spec.get("bold") is not None:
        font_proxy.bold = spec["bold"]
    if spec.get("italic") is not None:
        font_proxy.italic = spec["italic"]
    rpr = ensure_child(xml_parent, "w:rPr")
    if spec.get("font_name") is not None or spec.get("east_asia_font") is not None:
        rfonts = ensure_child(rpr, "w:rFonts")
        if spec.get("font_name") is not None:
            rfonts.set(qn("w:ascii"), spec["font_name"])
            rfonts.set(qn("w:hAnsi"), spec["font_name"])
            rfonts.set(qn("w:cs"), spec["font_name"])
        if spec.get("east_asia_font") is not None:
            rfonts.set(qn("w:eastAsia"), spec["east_asia_font"])
    if spec.get("size_pt") is not None:
        size_value = pt_to_half_points(spec["size_pt"])
        set_decimal_value(rpr, "w:sz", size_value)
        set_decimal_value(rpr, "w:szCs", size_value)
    set_on_off(rpr, "w:b", spec.get("bold"))
    set_on_off(rpr, "w:i", spec.get("italic"))
    if spec.get("color") is not None:
        color = ensure_child(rpr, "w:color")
        color.set(qn("w:val"), str(spec["color"]))


def alignment_value(value: str | None) -> Any:
    if value is None:
        return None
    return ALIGNMENT_MAP.get(value.lower())


def apply_paragraph_spec(paragraph_format: Any, spec: dict[str, Any], xml_parent: Any | None = None) -> None:
    alignment = alignment_value(spec.get("alignment"))
    if alignment is not None:
        paragraph_format.alignment = alignment
    if spec.get("space_before_pt") is not None:
        paragraph_format.space_before = Pt(spec["space_before_pt"])
    if spec.get("space_after_pt") is not None:
        paragraph_format.space_after = Pt(spec["space_after_pt"])
    if spec.get("line_spacing_pt") is not None:
        paragraph_format.line_spacing = Pt(spec["line_spacing_pt"])
    if spec.get("left_indent_pt") is not None:
        paragraph_format.left_indent = Pt(spec["left_indent_pt"])
    if spec.get("right_indent_pt") is not None:
        paragraph_format.right_indent = Pt(spec["right_indent_pt"])
    if spec.get("first_line_indent_pt") is not None:
        paragraph_format.first_line_indent = Pt(spec["first_line_indent_pt"])
    elif spec.get("hanging_indent_pt") is not None:
        paragraph_format.first_line_indent = Pt(-spec["hanging_indent_pt"])
    if spec.get("keep_lines") is not None:
        paragraph_format.keep_together = spec["keep_lines"]
    if spec.get("keep_next") is not None:
        paragraph_format.keep_with_next = spec["keep_next"]
    if spec.get("page_break_before") is not None:
        paragraph_format.page_break_before = spec["page_break_before"]
    if spec.get("widow_control") is not None:
        paragraph_format.widow_control = spec["widow_control"]
    if xml_parent is None:
        return
    ppr = ensure_child(xml_parent, "w:pPr")
    if spec.get("alignment") is not None:
        alignment_xml = ensure_child(ppr, "w:jc")
        alignment_xml.set(qn("w:val"), str(spec["alignment"]))
    if any(
        spec.get(field) is not None
        for field in ["space_before_pt", "space_after_pt", "line_spacing_pt", "line_spacing_rule"]
    ):
        spacing = ensure_child(ppr, "w:spacing")
        if spec.get("space_before_pt") is not None:
            spacing.set(qn("w:before"), pt_to_twips(spec["space_before_pt"]))
        if spec.get("space_after_pt") is not None:
            spacing.set(qn("w:after"), pt_to_twips(spec["space_after_pt"]))
        if spec.get("line_spacing_pt") is not None:
            spacing.set(qn("w:line"), pt_to_twips(spec["line_spacing_pt"]))
        if spec.get("line_spacing_rule") is not None:
            spacing.set(qn("w:lineRule"), str(spec["line_spacing_rule"]))
    if any(
        spec.get(field) is not None
        for field in ["left_indent_pt", "right_indent_pt", "first_line_indent_pt", "hanging_indent_pt"]
    ):
        indent = ensure_child(ppr, "w:ind")
        if spec.get("left_indent_pt") is not None:
            indent.set(qn("w:left"), pt_to_twips(spec["left_indent_pt"]))
        if spec.get("right_indent_pt") is not None:
            indent.set(qn("w:right"), pt_to_twips(spec["right_indent_pt"]))
        if spec.get("first_line_indent_pt") is not None:
            indent.set(qn("w:firstLine"), pt_to_twips(spec["first_line_indent_pt"]))
        if spec.get("hanging_indent_pt") is not None:
            indent.set(qn("w:hanging"), pt_to_twips(spec["hanging_indent_pt"]))
    set_on_off(ppr, "w:keepLines", spec.get("keep_lines"))
    set_on_off(ppr, "w:keepNext", spec.get("keep_next"))
    set_on_off(ppr, "w:pageBreakBefore", spec.get("page_break_before"))
    set_on_off(ppr, "w:widowControl", spec.get("widow_control"))
    if spec.get("outline_level") is not None:
        outline = ensure_child(ppr, "w:outlineLvl")
        outline.set(qn("w:val"), str(spec["outline_level"]))


def find_style(document: Document, rule: dict[str, Any]) -> Any | None:
    style_id = rule.get("style_id")
    style_name = rule.get("name")
    # Style IDs are not stable across unrelated DOCX files. Prefer the human-readable
    # name first so we do not accidentally rename `Normal` into `Heading 1`, etc.
    for style in document.styles:
        if style_name and style.name == style_name:
            return style
    for style in document.styles:
        if style_id and style.style_id == style_id:
            if not style_name or style.name == style_name:
                return style
    return None


def ensure_style(document: Document, rule: dict[str, Any]) -> Any | None:
    style = find_style(document, rule)
    if style is not None:
        return style
    style_name = rule.get("name")
    style_type = STYLE_TYPE_MAP.get(rule.get("style_type") or "paragraph")
    if not style_name or style_type is None:
        return None
    return document.styles.add_style(style_name, style_type)


def apply_style_rule(document: Document, rule: dict[str, Any]) -> bool:
    style = ensure_style(document, rule)
    if style is None:
        return False
    if rule.get("style_id") is not None:
        style.element.set(qn("w:styleId"), str(rule["style_id"]))
    if rule.get("style_type") is not None:
        style.element.set(qn("w:type"), str(rule["style_type"]))
    if rule.get("name") is not None:
        style_name = ensure_child(style.element, "w:name")
        style_name.set(qn("w:val"), str(rule["name"]))
    if rule.get("based_on") is not None:
        based_on = ensure_child(style.element, "w:basedOn")
        based_on.set(qn("w:val"), str(rule["based_on"]))
    if rule.get("next_style") is not None:
        next_style = ensure_child(style.element, "w:next")
        next_style.set(qn("w:val"), str(rule["next_style"]))
    apply_run_spec(style.font, style.element, rule.get("run", {}))
    if hasattr(style, "paragraph_format"):
        apply_paragraph_spec(style.paragraph_format, rule.get("paragraph", {}), style.element)
    return True


def apply_rule_to_paragraph(document: Document, paragraph: Any, rule: dict[str, Any]) -> None:
    style_name = rule.get("style_name")
    if style_name:
        target_style = None
        for style in document.styles:
            if style.name == style_name:
                target_style = style
                break
        if target_style is not None:
            paragraph.style = target_style
    apply_paragraph_spec(paragraph.paragraph_format, rule.get("paragraph", {}), paragraph._p)
    non_empty_runs = [run for run in paragraph.runs if run.text.strip()]
    for run in non_empty_runs:
        apply_run_spec(run.font, run._r, rule.get("run", {}))


def iter_document_header_paragraphs(document: Document, header_kind: str) -> list[Any]:
    attr_name = {
        "default": "header",
        "first_page": "first_page_header",
        "even_page": "even_page_header",
    }.get(header_kind, "header")
    paragraphs: list[Any] = []
    for section in document.sections:
        header = getattr(section, attr_name, None)
        if header is None:
            continue
        targets = [paragraph for paragraph in header.paragraphs if paragraph.text.strip()]
        if not targets and header.paragraphs:
            targets = [header.paragraphs[0]]
        paragraphs.extend(targets)
    return paragraphs


def iter_document_table_paragraphs(document: Document) -> list[Any]:
    paragraphs: list[Any] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph)
    return paragraphs


def apply_sections(document: Document, rules: dict[str, Any]) -> int:
    section_rule = (rules.get("section_settings") or [{}])[0]
    updated = 0
    for section in document.sections:
        orientation = section_rule.get("orientation")
        if orientation in ORIENTATION_MAP:
            section.orientation = ORIENTATION_MAP[orientation]
        for attr, field in [
            ("page_width", "page_width_pt"),
            ("page_height", "page_height_pt"),
            ("left_margin", "left_margin_pt"),
            ("right_margin", "right_margin_pt"),
            ("top_margin", "top_margin_pt"),
            ("bottom_margin", "bottom_margin_pt"),
            ("header_distance", "header_distance_pt"),
            ("footer_distance", "footer_distance_pt"),
        ]:
            value = section_rule.get(field)
            if value is not None:
                setattr(section, attr, Pt(value))
        updated += 1
    return updated


def apply_doc_defaults(document: Document, defaults: dict[str, Any]) -> None:
    styles_root = document.styles.element
    doc_defaults = ensure_child(styles_root, "w:docDefaults")
    run_default = ensure_child(doc_defaults, "w:rPrDefault")
    run_props = ensure_child(run_default, "w:rPr")
    para_default = ensure_child(doc_defaults, "w:pPrDefault")
    para_props = ensure_child(para_default, "w:pPr")

    run_spec = defaults.get("run", {})
    if run_spec.get("font_name") is not None or run_spec.get("east_asia_font") is not None:
        rfonts = ensure_child(run_props, "w:rFonts")
        if run_spec.get("font_name") is not None:
            rfonts.set(qn("w:ascii"), run_spec["font_name"])
            rfonts.set(qn("w:hAnsi"), run_spec["font_name"])
            rfonts.set(qn("w:cs"), run_spec["font_name"])
        if run_spec.get("east_asia_font") is not None:
            rfonts.set(qn("w:eastAsia"), run_spec["east_asia_font"])
    if run_spec.get("size_pt") is not None:
        size_value = pt_to_half_points(run_spec["size_pt"])
        set_decimal_value(run_props, "w:sz", size_value)
        set_decimal_value(run_props, "w:szCs", size_value)
    set_on_off(run_props, "w:b", run_spec.get("bold"))
    set_on_off(run_props, "w:i", run_spec.get("italic"))

    paragraph_spec = defaults.get("paragraph", {})
    apply_paragraph_spec(document.styles["Normal"].paragraph_format, paragraph_spec, para_default)


def clear_document_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def clear_paragraph_content(paragraph: Any) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def find_style_by_name(document: Document, style_name: str | None) -> Any | None:
    if not style_name:
        return None
    for style in document.styles:
        if style.name == style_name:
            return style
    return None


def apply_style_relationships(document: Document, tracked_styles: list[dict[str, Any]]) -> None:
    for rule in tracked_styles:
        style = find_style(document, rule)
        if style is None:
            continue
        based_on_name = rule.get("based_on")
        if based_on_name:
            based_on_style = find_style_by_name(document, based_on_name)
            if based_on_style is not None:
                try:
                    style.base_style = based_on_style
                except ValueError:
                    pass
        next_style_name = rule.get("next_style")
        if next_style_name and hasattr(style, "next_paragraph_style"):
            next_style = find_style_by_name(document, next_style_name)
            if next_style is not None:
                try:
                    style.next_paragraph_style = next_style
                except ValueError:
                    pass


def rebuild_paragraph_from_blueprint(document: Document, paragraph: Any, blueprint: dict[str, Any]) -> None:
    style = find_style_by_name(document, blueprint.get("style_name"))
    if style is not None:
        paragraph.style = style
    apply_paragraph_spec(paragraph.paragraph_format, blueprint.get("paragraph", {}), paragraph._p)

    clear_paragraph_content(paragraph)
    runs = blueprint.get("runs") or []
    if not runs and blueprint.get("text"):
        runs = [
            {
                "text": blueprint["text"],
                "style_name": None,
                "run": blueprint.get("run", {}),
            }
        ]
    for run_blueprint in runs:
        run = paragraph.add_run(run_blueprint.get("text") or "")
        run_style = find_style_by_name(document, run_blueprint.get("style_name"))
        if run_style is not None:
            run.style = run_style
        apply_run_spec(run.font, run._r, run_blueprint.get("run", {}))


def apply_section_blueprint(document: Document, blueprint: dict[str, Any]) -> None:
    body = document._body._element
    paragraphs = [child for child in body if child.tag == qn("w:p")]
    for section_break in blueprint.get("section_breaks", []):
        paragraph_index = section_break.get("after_paragraph_index")
        sect_pr_xml = section_break.get("sectPr_xml")
        if paragraph_index is None or sect_pr_xml is None:
            continue
        if not (0 <= int(paragraph_index) < len(paragraphs)):
            continue
        paragraph = paragraphs[int(paragraph_index)]
        ppr = ensure_child(paragraph, "w:pPr")
        existing = ppr.find(qn("w:sectPr"))
        if existing is not None:
            ppr.remove(existing)
        ppr.append(etree.fromstring(str(sect_pr_xml).encode("utf-8")))

    final_section_xml = blueprint.get("final_section_xml")
    if final_section_xml is None:
        return
    existing_final = body.find(qn("w:sectPr"))
    if existing_final is not None:
        body.remove(existing_final)
    body.append(etree.fromstring(str(final_section_xml).encode("utf-8")))


def rebuild_template_from_rule_pack(rule_pack_path: Path, output_path: Path) -> None:
    rebuild_docx_from_package_snapshot(package_snapshot_dir(rule_pack_path), output_path)


def compare_template_models(template_model: dict[str, Any], rebuilt_model: dict[str, Any]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []

    compare_mapping(
        findings,
        expected=template_model.get("document_defaults", {}).get("paragraph", {}),
        actual=rebuilt_model.get("document_defaults", {}).get("paragraph", {}),
        category="template-default-paragraph",
        location="document defaults paragraph",
        severity="error",
        fields=[
            "alignment",
            "space_before_pt",
            "space_after_pt",
            "line_spacing_pt",
            "first_line_indent_pt",
        ],
        tolerance_map={
            "space_before_pt": 1.0,
            "space_after_pt": 1.0,
            "line_spacing_pt": 1.0,
            "first_line_indent_pt": 1.0,
        },
    )
    compare_mapping(
        findings,
        expected=template_model.get("document_defaults", {}).get("run", {}),
        actual=rebuilt_model.get("document_defaults", {}).get("run", {}),
        category="template-default-run",
        location="document defaults run",
        severity="error",
        fields=["font_name", "east_asia_font", "size_pt", "bold", "italic"],
        tolerance_map={"size_pt": 0.5},
    )

    expected_sections = template_model.get("sections", [])
    actual_sections = rebuilt_model.get("sections", [])
    if len(expected_sections) != len(actual_sections):
        add_finding(
            findings,
            severity="error",
            category="template-structure",
            location="document",
            field="section_count",
            expected=len(expected_sections),
            actual=len(actual_sections),
            message="Section count does not match the original template.",
        )
    for index, expected_section in enumerate(expected_sections[: len(actual_sections)]):
        actual_section = actual_sections[index]
        compare_mapping(
            findings,
            expected=expected_section,
            actual=actual_section,
            category="template-section",
            location=f"template section {index}",
            severity="error",
            fields=[
                "orientation",
                "page_width_pt",
                "page_height_pt",
                "left_margin_pt",
                "right_margin_pt",
                "top_margin_pt",
                "bottom_margin_pt",
                "header_distance_pt",
                "footer_distance_pt",
            ],
            tolerance_map={
                "page_width_pt": 1.0,
                "page_height_pt": 1.0,
                "left_margin_pt": 1.0,
                "right_margin_pt": 1.0,
                "top_margin_pt": 1.0,
                "bottom_margin_pt": 1.0,
                "header_distance_pt": 1.0,
                "footer_distance_pt": 1.0,
            },
        )

    expected_paragraphs = template_model.get("paragraphs", [])
    actual_paragraphs = rebuilt_model.get("paragraphs", [])
    if len(expected_paragraphs) != len(actual_paragraphs):
        add_finding(
            findings,
            severity="error",
            category="template-structure",
            location="document",
            field="paragraph_count",
            expected=len(expected_paragraphs),
            actual=len(actual_paragraphs),
            message="Paragraph count does not match the original template.",
        )

    paragraph_fields = [
        "alignment",
        "space_before_pt",
        "space_after_pt",
        "line_spacing_pt",
        "first_line_indent_pt",
        "keep_next",
        "keep_lines",
        "page_break_before",
        "widow_control",
        "outline_level",
    ]
    for index, expected_paragraph in enumerate(expected_paragraphs[: len(actual_paragraphs)]):
        actual_paragraph = actual_paragraphs[index]
        location = f"paragraph {index}"
        compare_mapping(
            findings,
            expected={
                "text": expected_paragraph.get("text"),
                "style_name": expected_paragraph.get("style_name"),
                "semantic_label": expected_paragraph.get("semantic_label"),
            },
            actual={
                "text": actual_paragraph.get("text"),
                "style_name": actual_paragraph.get("style_name"),
                "semantic_label": actual_paragraph.get("semantic_label"),
            },
            category="template-paragraph",
            location=location,
            severity="error",
            fields=["text", "style_name", "semantic_label"],
        )
        compare_mapping(
            findings,
            expected=expected_paragraph.get("paragraph", {}),
            actual=actual_paragraph.get("paragraph", {}),
            category="template-paragraph",
            location=location,
            severity="error",
            fields=paragraph_fields,
            tolerance_map={
                "space_before_pt": 1.0,
                "space_after_pt": 1.0,
                "line_spacing_pt": 1.0,
                "first_line_indent_pt": 1.0,
            },
        )

        expected_runs = expected_paragraph.get("runs", [])
        actual_runs = actual_paragraph.get("runs", [])
        if len(expected_runs) != len(actual_runs):
            add_finding(
                findings,
                severity="error",
                category="template-run",
                location=location,
                field="run_count",
                expected=len(expected_runs),
                actual=len(actual_runs),
                message="Run count does not match the original template paragraph.",
            )
        for run_index, expected_run in enumerate(expected_runs[: len(actual_runs)]):
            actual_run = actual_runs[run_index]
            run_location = f"paragraph {index} / run {run_index}"
            compare_mapping(
                findings,
                expected={
                    "text": expected_run.get("text"),
                    "style_name": expected_run.get("style_name"),
                },
                actual={
                    "text": actual_run.get("text"),
                    "style_name": actual_run.get("style_name"),
                },
                category="template-run",
                location=run_location,
                severity="error",
                fields=["text", "style_name"],
            )
            compare_mapping(
                findings,
                expected=expected_run.get("run", {}),
                actual=actual_run.get("run", {}),
                category="template-run",
                location=run_location,
                severity="error",
                fields=["font_name", "east_asia_font", "size_pt", "bold", "italic"],
                tolerance_map={"size_pt": 0.5},
            )

    actual_styles = {style["style_id"]: style for style in rebuilt_model.get("styles", {}).values()}
    actual_style_names = {
        style["name"]: style for style in rebuilt_model.get("styles", {}).values() if style.get("name")
    }
    for expected_style in template_model.get("tracked_styles", []):
        actual_style = actual_styles.get(expected_style.get("style_id")) or actual_style_names.get(
            expected_style.get("name")
        )
        location = expected_style.get("name") or expected_style.get("style_id") or "style"
        if actual_style is None:
            add_finding(
                findings,
                severity="error",
                category="template-style",
                location=location,
                field="exists",
                expected="present",
                actual="missing",
                message="Tracked style is missing from the rebuilt template.",
            )
            continue
        compare_mapping(
            findings,
            expected={
                "based_on": expected_style.get("based_on"),
                "next_style": expected_style.get("next_style"),
            },
            actual={
                "based_on": actual_style.get("based_on"),
                "next_style": actual_style.get("next_style"),
            },
            category="template-style",
            location=location,
            severity="error",
            fields=["based_on", "next_style"],
        )
        compare_mapping(
            findings,
            expected=expected_style.get("paragraph", {}),
            actual=actual_style.get("paragraph", {}),
            category="template-style",
            location=location,
            severity="error",
            fields=["alignment", "space_before_pt", "space_after_pt", "line_spacing_pt", "outline_level"],
            tolerance_map={"space_before_pt": 1.0, "space_after_pt": 1.0, "line_spacing_pt": 1.0},
        )
        compare_mapping(
            findings,
            expected=expected_style.get("run", {}),
            actual=actual_style.get("run", {}),
            category="template-style",
            location=location,
            severity="error",
            fields=["font_name", "east_asia_font", "size_pt", "bold", "italic"],
            tolerance_map={"size_pt": 0.5},
        )

    return findings


def render_validation_markdown(report: dict[str, Any], root: Path) -> str:
    text_requirement_summary = report.get("text_requirement_summary") or {}
    lines = [
        "# Rule Pack Validation",
        "",
        f"- Generated at: `{report['generated_at']}`",
        f"- Rule pack: `{relative_path(Path(report['rule_pack_path']), root)}`",
        f"- Template: `{relative_path(Path(report['template_path']), root)}`",
        f"- Rebuilt template: `{relative_path(Path(report['rebuilt_template_path']), root)}`",
        f"- Comparison level: `docx package metadata`",
        f"- Rebuild source: `{report.get('rebuild_source', 'persisted-rule-pack-payload')}`",
        f"- Valid: `{'yes' if report['summary']['pass'] else 'no'}`",
        f"- Errors: `{report['summary']['errors']}`",
        f"- Warnings: `{report['summary']['warnings']}`",
        f"- Prose-derived rules: `{text_requirement_summary.get('supported', 0)}` active / `{text_requirement_summary.get('total', 0)}` total",
        "",
        "## Findings",
        "",
    ]
    if not report["findings"]:
        lines.extend(["No DOCX package metadata differences were detected between the rebuilt template and the original template.", ""])
        return "\n".join(lines)
    for finding in report["findings"]:
        lines.append(
            "- [{severity}] {category} | {location} | {field} | expected `{expected}` | actual `{actual}` | {message}".format(
                severity=finding["severity"].upper(),
                category=finding["category"],
                location=finding["location"],
                field=finding["field"],
                expected=format_report_value(finding["expected"]),
                actual=format_report_value(finding["actual"]),
                message=finding["message"],
            )
        )
        context = finding.get("context") or {}
        if context.get("paragraph_index") is not None:
            lines.append(
                "  Paragraph `{index}` snippet: first 18 chars `{prefix}` | last 18 chars `{suffix}`".format(
                    index=context["paragraph_index"],
                    prefix=context.get("text_prefix", ""),
                    suffix=context.get("text_suffix", ""),
                )
            )
    lines.append("")
    return "\n".join(lines)


def validate_rule_pack_contents(
    *,
    root: Path,
    rule_pack_path: Path,
    pack: dict[str, Any],
    template_path: Path,
    template_model: dict[str, Any] | None = None,
) -> dict[str, Any]:
    artifact_dir = validation_artifact_dir(rule_pack_path)
    rebuilt_template_path = artifact_dir / "rebuilt-template.docx"
    report_json = artifact_dir / "validation-report.json"
    report_md = artifact_dir / "validation-report.md"

    template_model = template_model or build_document_model(template_path)
    try:
        rebuild_template_from_rule_pack(rule_pack_path, rebuilt_template_path)
        expected_manifest = read_json(package_manifest_path(rule_pack_path), default=None)
        if expected_manifest is None:
            expected_manifest = extract_docx_package_snapshot(template_path, package_snapshot_dir(rule_pack_path))
            write_json(package_manifest_path(rule_pack_path), expected_manifest)
        actual_manifest = manifest_from_docx_package(rebuilt_template_path)
        findings = compare_package_manifests(expected_manifest, actual_manifest)
    except Exception as exc:
        findings = [
            {
                "severity": "error",
                "category": "docx-package",
                "location": "document",
                "field": "rebuild",
                "expected": "no package metadata differences",
                "actual": str(exc),
                "message": "DOCX package validation failed before diffing.",
            }
        ]

    summary = {
        "errors": sum(1 for item in findings if item["severity"] == "error"),
        "warnings": sum(1 for item in findings if item["severity"] == "warning"),
        "info": sum(1 for item in findings if item["severity"] == "info"),
    }
    summary["pass"] = summary["errors"] == 0

    report = {
        "version": RULE_PACK_VERSION,
        "generated_at": utc_now(),
        "pipeline_version": PIPELINE_VERSION,
        "rule_pack_path": str(rule_pack_path.resolve()),
        "template_path": str(template_path.resolve()),
        "rebuilt_template_path": str(rebuilt_template_path.resolve()),
        "comparison": "docx-package-metadata",
        "package_manifest": str(package_manifest_path(rule_pack_path).resolve()),
        "rebuild_source": "persisted-rule-pack-payload",
        "text_requirement_summary": summarize_text_requirement_rules(pack["rules"].get("text_requirement_rules", [])),
        "summary": summary,
        "findings": enrich_findings_with_context(findings, template_model),
    }
    write_json(report_json, report)
    write_text(report_md, render_validation_markdown(report, root))
    report["report_json"] = str(report_json.resolve())
    report["report_md"] = str(report_md.resolve())
    return report


def repair_document(args: argparse.Namespace, root: Path) -> int:
    document_path = Path(args.document).resolve()
    if not document_path.exists():
        raise FileNotFoundError(f"Thesis DOCX not found: {document_path}")
    rule_pack_path = find_rule_pack_path(root, args.rule_pack)
    if rule_pack_path is None:
        raise FileNotFoundError("No rule pack found. Build one first or pass --rule-pack.")
    pack = load_rule_pack(rule_pack_path)
    ensure_validated_rule_pack(rule_pack_path, pack)
    registry = load_registry(root)
    matched_audit = find_matching_audit(registry, document_path, rule_pack_path)
    if matched_audit is None:
        raise RuntimeError("Repair is blocked until the target thesis has been audited with the same rule pack.")

    output_path = (
        Path(args.output).resolve()
        if args.output
        else root / "artifacts" / "repaired" / f"{path_slug(document_path, root)}-formatted.docx"
    )
    ensure_dir(output_path.parent)

    rules = pack["rules"]
    document = Document(str(document_path))
    original_model = build_document_model(document_path)

    style_updates = sum(1 for rule in rules.get("tracked_styles", []) if apply_style_rule(document, rule))
    section_updates = apply_sections(document, rules)

    semantic_targets = model_by_semantic_label(original_model)
    paragraph_updates = 0
    for label, rule in rules.get("semantic_rules", {}).items():
        actual = semantic_targets.get(label)
        if actual is None:
            continue
        apply_rule_to_paragraph(document, document.paragraphs[actual["index"]], rule)
        paragraph_updates += 1

    for index, rule in enumerate(rules.get("cover_rules", {}).get("blocks", [])):
        if index >= len(original_model.get("cover_blocks", [])):
            break
        actual = original_model["cover_blocks"][index]
        apply_rule_to_paragraph(document, document.paragraphs[actual["index"]], rule)
        paragraph_updates += 1

    for custom_rule in rules.get("custom_paragraph_rules", []):
        for paragraph in paragraphs_for_custom_rule(original_model, custom_rule):
            apply_rule_to_paragraph(document, document.paragraphs[paragraph["index"]], custom_rule)
            paragraph_updates += 1

    for header_rule in rules.get("header_rules", {}).values():
        for paragraph in iter_document_header_paragraphs(document, header_rule.get("header_kind", "default")):
            apply_rule_to_paragraph(document, paragraph, header_rule)
            paragraph_updates += 1

    for table_rule in rules.get("custom_table_rules", []):
        for paragraph in iter_document_table_paragraphs(document):
            apply_rule_to_paragraph(document, paragraph, table_rule)
            paragraph_updates += 1

    document.save(str(output_path))

    repaired_model = build_document_model(output_path)
    post_audit = audit_against_rule_pack(repaired_model, pack)
    report_dir = (
        Path(args.report_dir).resolve()
        if args.report_dir
        else root / "artifacts" / "repairs" / path_slug(document_path, root)
    )
    ensure_dir(report_dir)
    summary_path = report_dir / "repair-summary.json"
    audit_json_path = report_dir / "post-repair-audit.json"
    audit_md_path = report_dir / "post-repair-audit.md"
    summary_payload = {
        "generated_at": utc_now(),
        "document_path": str(document_path.resolve()),
        "output_path": str(output_path.resolve()),
        "rule_pack_path": str(rule_pack_path.resolve()),
        "style_updates": style_updates,
        "section_updates": section_updates,
        "paragraph_updates": paragraph_updates,
    }
    report = {
        "version": RULE_PACK_VERSION,
        "generated_at": utc_now(),
        "pipeline_version": PIPELINE_VERSION,
        "rule_pack_path": str(rule_pack_path.resolve()),
        "document_path": str(output_path.resolve()),
        "document_sha256": repaired_model["source"]["sha256"],
        "summary": post_audit["summary"],
        "findings": enrich_findings_with_context_v2(post_audit["findings"], repaired_model),
    }
    write_json(summary_path, summary_payload)
    write_json(audit_json_path, report)
    write_text(audit_md_path, render_audit_markdown_v2(report, root))

    repair_entry = {
        "generated_at": summary_payload["generated_at"],
        "document_path": str(document_path.resolve()),
        "output_path": str(output_path.resolve()),
        "rule_pack_path": str(rule_pack_path.resolve()),
        "summary_path": str(summary_path.resolve()),
        "post_audit_json": str(audit_json_path.resolve()),
        "post_audit_md": str(audit_md_path.resolve()),
        "summary": report["summary"],
    }
    registry["repairs"].append(repair_entry)
    registry["latest_repair"] = repair_entry
    save_registry(root, registry)

    print(
        json.dumps(
            {
                "status": "ok",
                "action": "repair",
                "output_path": str(output_path.resolve()),
                "summary_path": str(summary_path.resolve()),
                "post_audit_json": str(audit_json_path.resolve()),
                "summary": report["summary"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


def validate_rule_pack_command(args: argparse.Namespace, root: Path) -> int:
    rule_pack_path = find_rule_pack_path(root, args.rule_pack)
    if rule_pack_path is None:
        raise FileNotFoundError("No rule pack found. Build one first or pass --rule-pack.")

    pack = load_rule_pack(rule_pack_path)
    manifest = dict(pack["manifest"])
    template_value = args.template or manifest.get("template", {}).get("path")
    if not template_value:
        raise FileNotFoundError("Template path is missing. Pass --template <template.docx> to validate the rule pack.")
    template_path = Path(template_value).resolve()
    if not template_path.exists():
        raise FileNotFoundError(f"Template DOCX not found: {template_path}")

    validation = validate_rule_pack_contents(
        root=root,
        rule_pack_path=rule_pack_path,
        pack=pack,
        template_path=template_path,
    )
    manifest["validation"] = {
        "validated_at": validation["generated_at"],
        "comparison": validation["comparison"],
        "pass": validation["summary"]["pass"],
        "errors": validation["summary"]["errors"],
        "warnings": validation["summary"]["warnings"],
        "rebuilt_template_path": validation["rebuilt_template_path"],
        "package_manifest": validation["package_manifest"],
        "report_json": validation["report_json"],
        "report_md": validation["report_md"],
    }
    write_json(rule_pack_path / "manifest.json", manifest)

    template_sha = manifest.get("template", {}).get("sha256")
    entry = {
        "name": manifest.get("name") or rule_pack_path.name,
        "path": str(rule_pack_path.resolve()),
        "template_path": str(template_path.resolve()),
        "template_sha256": template_sha,
        "generated_at": manifest.get("generated_at"),
        "validation_comparison": validation["comparison"],
        "validation_pass": validation["summary"]["pass"],
        "validation_report_json": validation["report_json"],
        "validation_report_md": validation["report_md"],
    }
    register_rule_pack(root, entry)

    print(
        json.dumps(
            {
                "status": "ok" if validation["summary"]["pass"] else "invalid",
                "action": "validate-rule-pack",
                "rule_pack_path": str(rule_pack_path.resolve()),
                "validation_report_json": validation["report_json"],
                "validation_report_md": validation["report_md"],
                "summary": validation["summary"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if validation["summary"]["pass"] else 1


def show_status(root: Path) -> int:
    print(json.dumps(load_registry(root), ensure_ascii=False, indent=2))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Template-driven thesis DOCX format workflow.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    build_parser_ = subparsers.add_parser("build-rule-pack", help="Build a persistent rule pack from a template DOCX.")
    build_parser_.add_argument("--template", required=True, help="Path to the thesis template DOCX.")
    build_parser_.add_argument("--name", help="Rule pack name. Defaults to the template stem.")
    build_parser_.add_argument("--output-dir", help="Output directory for the rule pack.")

    validate_parser = subparsers.add_parser(
        "validate-rule-pack",
        help="Rebuild the template from a rule pack and compare it against the original template.",
    )
    validate_parser.add_argument("--rule-pack", required=True, help="Rule pack directory to validate.")
    validate_parser.add_argument("--template", help="Template DOCX path. Defaults to the path stored in the manifest.")

    audit_parser = subparsers.add_parser("audit", help="Audit a thesis DOCX against a rule pack.")
    audit_parser.add_argument("--document", required=True, help="Path to the thesis DOCX.")
    audit_parser.add_argument("--rule-pack", help="Rule pack directory. Defaults to the latest built pack.")
    audit_parser.add_argument("--report-dir", help="Output directory for audit artifacts.")

    repair_parser = subparsers.add_parser("repair", help="Repair a thesis DOCX after it has been audited.")
    repair_parser.add_argument("--document", required=True, help="Path to the thesis DOCX.")
    repair_parser.add_argument("--rule-pack", help="Rule pack directory. Defaults to the latest built pack.")
    repair_parser.add_argument("--output", help="Output path for the repaired DOCX.")
    repair_parser.add_argument("--report-dir", help="Output directory for repair artifacts.")
    repair_parser.add_argument(
        "--approval-id",
        help="One-time user approval token issued by the project hooks after audit review.",
    )

    subparsers.add_parser("status", help="Show the workflow registry.")
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    root = workspace_root(Path.cwd())
    try:
        if args.command == "build-rule-pack":
            return build_rule_pack(args, root)
        if args.command == "validate-rule-pack":
            return validate_rule_pack_command(args, root)
        if args.command == "audit":
            return audit_document(args, root)
        if args.command == "repair":
            return repair_document(args, root)
        if args.command == "status":
            return show_status(root)
        parser.error(f"Unknown command: {args.command}")
    except Exception as exc:  # pragma: no cover
        print(json.dumps({"status": "error", "message": str(exc)}, ensure_ascii=False, indent=2))
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
